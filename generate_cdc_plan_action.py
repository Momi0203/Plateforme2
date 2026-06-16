"""
Génération du Cahier des Charges — Module Plan d'Action de la Plateforme SIG
Exécuter depuis n'importe où : python generate_cdc_plan_action.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "CDC_Module_PlanAction_PlateformeSIG.docx")

# ── Couleurs ─────────────────────────────────────────────────────────────────
C_BLEU_TITRE  = RGBColor(0x1F, 0x49, 0x7D)
C_BLEU_H2     = RGBColor(0x2E, 0x74, 0xB5)
C_BLEU_H3     = RGBColor(0x2F, 0x5A, 0x8C)
C_ORANGE      = RGBColor(0xC5, 0x5A, 0x11)
C_VERT        = RGBColor(0x37, 0x86, 0x30)
C_GRIS_CELL   = RGBColor(0xBF, 0xD7, 0xED)
C_BLEU_FONCE  = RGBColor(0x1F, 0x49, 0x7D)
C_BLANC       = RGBColor(0xFF, 0xFF, 0xFF)
C_GRIS_ALT    = RGBColor(0xF2, 0xF2, 0xF2)
C_VERT_CELL   = RGBColor(0xE2, 0xEF, 0xDA)
C_ROUGE_CELL  = RGBColor(0xFF, 0xE0, 0xE0)


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '4472C4')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = C_BLEU_TITRE
    run.font.size = Pt(16)
    run.bold = True
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = C_BLEU_H2
    run.font.size = Pt(13)
    run.bold = True
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = C_BLEU_H3
    run.font.size = Pt(11)
    run.bold = True
    return p


def body(doc, text, indent=0):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("Note : ")
    run.bold = True
    run.font.color.rgb = C_ORANGE
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = C_ORANGE
    p.paragraph_format.space_after = Pt(6)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_BLEU_FONCE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_BLANC
        run.font.size = Pt(10)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = C_GRIS_ALT if ri % 2 == 0 else C_BLANC
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = Cm(col_widths[ci])

    set_cell_borders(table)
    doc.add_paragraph()
    return table


def make_rights_table(doc, headers, rows, col_widths=None):
    """Table de droits avec coloration ✓/✗."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_BLEU_FONCE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_BLANC
        run.font.size = Pt(10)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = C_GRIS_ALT if ri % 2 == 0 else C_BLANC
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            if cell_text in ("✓", "Oui"):
                set_cell_bg(cell, C_VERT_CELL)
            elif cell_text in ("✗", "Non"):
                set_cell_bg(cell, C_ROUGE_CELL)
            elif ci == 0:
                set_cell_bg(cell, bg)
            else:
                set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            if cell_text == "✓":
                run.font.color.rgb = C_VERT
                run.bold = True
            elif cell_text == "✗":
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.bold = True

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = Cm(col_widths[ci])

    set_cell_borders(table)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("PLATEFORME SIG — GESTION DES RESSOURCES EN EAU")
run.font.size = Pt(14)
run.font.color.rgb = C_BLEU_TITRE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Périmètres irrigués — Tafilalet / Midelt")
run.font.size = Pt(11)
run.font.color.rgb = C_BLEU_H2

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CAHIER DES CHARGES")
run.font.size = Pt(28)
run.font.color.rgb = C_BLEU_TITRE
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MODULE « PLAN D\'ACTION »')
run.font.size = Pt(22)
run.font.color.rgb = C_ORANGE
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plans d'Aménagement  ·  Calendriers d'Intervention  ·  Suivi d'Avancement")
run.font.size = Pt(11)
run.font.color.rgb = C_BLEU_H3
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Version",          "1.0"),
    ("Date",             "Juin 2026"),
    ("Statut",           "Projet — en attente de validation"),
    ("Plateforme",       "Django 6 / GeoDjango / PostGIS"),
    ("Environnement",    "Windows 11 — OSGeo4W — PostgreSQL 15"),
    ("Auteur",           "Équipe de développement PFE"),
]
table = doc.add_table(rows=len(meta), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta):
    row = table.rows[i]
    c0 = row.cells[0]
    c1 = row.cells[1]
    set_cell_bg(c0, C_GRIS_CELL)
    set_cell_bg(c1, C_BLANC)
    run_k = c0.paragraphs[0].add_run(k)
    run_k.bold = True
    run_k.font.size = Pt(10)
    run_v = c1.paragraphs[0].add_run(v)
    run_v.font.size = Pt(10)
    c0.width = Cm(5)
    c1.width = Cm(8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SOMMAIRE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "Table des matières")
toc_items = [
    ("1.",    "Présentation générale et contexte"),
    ("2.",    "Objectifs du module"),
    ("3.",    "Référentiel des types d'action PMH"),
    ("4.",    "Axe 1 — Plans d'aménagement"),
    ("  4.1", "Modèle de données"),
    ("  4.2", "Fonctionnalités"),
    ("5.",    "Axe 2 — Calendrier d'intervention"),
    ("  5.1", "Modèle de données"),
    ("  5.2", "Fonctionnalités — Formulaire et Gantt"),
    ("  5.3", "Vue réseau PERT"),
    ("6.",    "Axe 3 — Suivi d'avancement"),
    ("  6.1", "Modèle de données"),
    ("  6.2", "Fonctionnalités"),
    ("7.",    "Matrice des droits"),
    ("8.",    "Architecture technique"),
    ("9.",    "Exigences non fonctionnelles"),
    ("10.",   "Critères d'acceptation"),
    ("11.",   "Livrables et planning indicatif"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 if num.startswith(" ") else 0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{num.strip()}  {title}")
    run.font.size = Pt(10.5)
    if not num.startswith(" "):
        run.bold = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  1. PRÉSENTATION GÉNÉRALE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "1. Présentation générale et contexte")

heading2(doc, "1.1 Contexte")
body(doc, (
    "Le module « Plan d'Action » est la sixième application fonctionnelle de la Plateforme SIG. "
    "Il centralise la planification, la programmation et le suivi opérationnel des travaux "
    "d'aménagement hydro-agricole relevant de la Petite et Moyenne Hydraulique (PMH) dans la "
    "zone d'action de l'ORMVA du Tafilalet (provinces de Midelt et Errachidia)."
))
body(doc, (
    "Dans le cadre de la stratégie Génération Green 2021-2030, l'ORMVA du Tafilalet gère "
    "annuellement un programme d'aménagement couvrant environ 230 km de réseaux d'irrigation "
    "traditionnels (séguias, khettaras), des seuils de dérivation, des ouvrages de protection "
    "et des interventions de modernisation. Ce module doit formaliser et tracer ce programme "
    "de manière structurée, collaborative et auditée."
))

heading2(doc, "1.2 Trois axes fonctionnels")
make_table(doc,
    ["Axe", "Intitulé", "Objet"],
    [
        ("A1", "Plans d'aménagement",
         "Enregistrement et consultation des plans annuels d'investissement PMH, "
         "classifiés par commune et par type d'action."),
        ("A2", "Calendrier d'intervention",
         "Décomposition de chaque action en tâches planifiées (Gantt / PERT), "
         "avec affectation des responsables et définition du mode de réalisation."),
        ("A3", "Suivi d'avancement",
         "Saisie périodique de l'état d'avancement par les acteurs identifiés "
         "dans le calendrier, avec dépôt de pièces justificatives."),
    ],
    [1.0, 4.5, 13.0]
)

heading2(doc, "1.3 Intégration dans la plateforme existante")
make_table(doc,
    ["Composant", "Technologie"],
    [
        ("Framework web",       "Django 6.0.4 — nouvelle app plan_action/"),
        ("Base de données",     "PostgreSQL 15 + PostGIS"),
        ("Gestion utilisateurs","compte.Utilisateur (AbstractUser + rôle : visiteur / opérateur / éditeur)"),
        ("Lien diagnostic",     "FK vers diagnostic.Perimetre et carte.Commune"),
        ("Gantt front-end",     "Frappe Gantt (MIT, CDN)"),
        ("Graphe PERT",         "vis.js Network (MIT/Apache, CDN)"),
        ("Courbe d'avancement", "Chart.js (MIT)"),
        ("Galerie photos",      "Lightbox2 (MIT)"),
        ("Export Excel",        "openpyxl (déjà présent dans le venv)"),
    ],
    [5.5, 13.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  2. OBJECTIFS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "2. Objectifs du module")

bullet(doc, "OBJ-1 : Fournir une vision consolidée des plans d'investissement PMH annuels "
            "par commune et par type d'action.")
bullet(doc, "OBJ-2 : Permettre la planification détaillée de chaque action sous forme de "
            "calendrier temporel (Gantt) et de réseau de tâches (PERT).")
bullet(doc, "OBJ-3 : Assurer le suivi multi-acteurs de l'avancement des travaux, "
            "études et procédures administratives.")
bullet(doc, "OBJ-4 : Tracer les pièces justificatives (PV d'attachement, PV de réception, "
            "photos de chantier, rapports d'étude) pour chaque étape.")
bullet(doc, "OBJ-5 : Appliquer une matrice de droits à quatre niveaux cohérente avec "
            "le modèle de rôles existant de la plateforme.")
bullet(doc, "OBJ-6 : Calculer automatiquement le chemin critique (CPM) et détecter "
            "les tâches en retard par rapport au calendrier prévu.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  3. RÉFÉRENTIEL DES TYPES D'ACTION PMH
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "3. Référentiel des types d'action PMH")

body(doc, (
    "Le référentiel ci-dessous est issu des programmes annuels de l'ORMVA du Tafilalet "
    "(Génération Green 2021-2030) et des interventions documentées sur les périmètres de PMH "
    "(séguias, khettaras, seuils, barrages collinaires). Il est implémenté sous forme de "
    "choices Django dans le modèle ActionPlan — liste fixe en V1, extensible par migration."
))

make_table(doc,
    ["Code", "Libellé", "Description / Unité de mesure"],
    [
        ("ACT-01", "Réhabilitation de séguias",
         "Revêtement, réfection de canaux d'irrigation en terre. Unité : ml (mètres linéaires)"),
        ("ACT-02", "Construction de séguias neuves",
         "Création de nouveaux canaux d'irrigation bétonnés ou en maçonnerie. Unité : ml"),
        ("ACT-03", "Construction de seuils de dérivation",
         "Ouvrages de prise d'eau sur oued (seuil déversant ou à vannes). Unité : nombre"),
        ("ACT-04", "Réhabilitation de khettaras",
         "Galeries souterraines d'irrigation traditionnelle — curage, soutènement. Unité : ml"),
        ("ACT-05", "Construction / réhabilitation de barrages collinaires",
         "Petites retenues d'eau pour irrigation gravitaire. Unité : nombre / m³ capacité"),
        ("ACT-06", "Aménagement de prises d'eau locales",
         "Prises sur source ou oued — maçonnerie, pertuis de régulation. Unité : nombre"),
        ("ACT-07", "Renforcement de murs de protection",
         "Protection des périmètres contre les crues et l'érosion de berge. Unité : ml"),
        ("ACT-08", "Entretien et curage de canaux",
         "Maintenance préventive des réseaux (désensablement, fauchage). Unité : ml"),
        ("ACT-09", "Réhabilitation de forages / puits",
         "Développement, équipement et remise en état de forages ou puits. Unité : nombre"),
        ("ACT-10", "Irrigation localisée (goutte à goutte)",
         "Modernisation — pose de réseau d'irrigation localisée. Unité : ha"),
        ("ACT-11", "Planage et nivellement de parcelles",
         "Mise en valeur agricole — terrassement, aménagement parcellaire. Unité : ha"),
        ("ACT-12", "Aménagement de pistes d'accès",
         "Infrastructure d'accès aux périmètres et ouvrages. Unité : ml"),
        ("ACT-13", "Protection contre les crues (digues, épis)",
         "Ouvrages de protection hydraulique — enrochements, épis, seuils de fond. Unité : ml"),
        ("ACT-14", "Étude technique préalable (APD / APS)",
         "Études d'avant-projet sommaire ou détaillé précédant les travaux. Unité : forfait"),
        ("ACT-15", "Autre",
         "Action hors nomenclature — préciser dans le champ description."),
    ],
    [1.8, 5.5, 11.2]
)

note(doc, "En V2, ce référentiel pourra être externalisé dans un modèle TypeAction "
          "administrable depuis /admin/, permettant l'ajout de nouvelles catégories "
          "sans migration Django.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  4. AXE 1 — PLANS D'AMÉNAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "4. Axe 1 — Plans d'aménagement")

# ─── 4.1 Modèle de données ─────────────────────────────────────────────────

heading2(doc, "4.1 Modèle de données")

heading3(doc, "PlanAmenagement — plan annuel global")
make_table(doc,
    ["Champ", "Type Django", "Contraintes / Choix"],
    [
        ("annee",              "PositiveIntegerField",
         "Unique. Plage 2000–2050. Représente l'exercice budgétaire annuel."),
        ("titre",              "CharField(200)",
         "Ex. : « Programme PMH 2026 — Province d'Errachidia »"),
        ("budget_total",       "DecimalField(15,2)",
         "Budget global en MAD (dirhams)."),
        ("source_financement", "CharField(choices)",
         "Budget État / FEADER / Collectivité territoriale / Partenariat / Autre"),
        ("statut",             "CharField(choices)",
         "en_preparation / publie / archive. Valeur par défaut : en_preparation."),
        ("description",        "TextField(blank=True)",
         "Contexte, objectifs généraux du plan."),
        ("date_creation",      "DateTimeField(auto_now_add=True)", "—"),
        ("cree_par",           "FK → Utilisateur(null=True)", "Enregistré automatiquement."),
    ],
    [4.0, 4.5, 10.0]
)

heading3(doc, "ActionPlan — ligne du plan")
make_table(doc,
    ["Champ", "Type Django", "Contraintes / Choix"],
    [
        ("plan",               "FK → PlanAmenagement", "CASCADE delete."),
        ("commune",            "FK → carte.Commune", "Localisation administrative de l'action."),
        ("perimetre",          "FK → diagnostic.Perimetre (null=True, blank=True)",
         "Optionnel — périmètre spécifique concerné."),
        ("type_action",        "CharField(choices)",
         "ACT-01 à ACT-15 (voir section 3)."),
        ("description",        "TextField",
         "Description détaillée de l'action."),
        ("budget_prevu",       "DecimalField(12,2)",
         "Budget prévisionnel en MAD."),
        ("superficie_concernee","DecimalField(8,2, null=True)",
         "Superficie en hectares (pour ACT-10, ACT-11)."),
        ("longueur_prevue",    "DecimalField(10,2, null=True)",
         "Longueur en mètres linéaires (pour ACT-01..04, 07, 08, 12, 13)."),
        ("statut",             "CharField(choices)",
         "programme / en_cours / realise / annule. Défaut : programme."),
        ("priorite",           "IntegerField(choices)",
         "1 = Haute / 2 = Moyenne / 3 = Basse."),
        ("observations",       "TextField(blank=True)", "—"),
        ("date_creation",      "DateTimeField(auto_now_add=True)", "—"),
        ("modifie_par",        "FK → Utilisateur(null=True)", "—"),
    ],
    [4.0, 4.5, 10.0]
)

# ─── 4.2 Fonctionnalités ───────────────────────────────────────────────────

heading2(doc, "4.2 Fonctionnalités")

make_table(doc,
    ["ID", "Fonctionnalité", "Priorité"],
    [
        ("F-A1-01", "Liste paginée des plans annuels avec indicateurs synthétiques "
                    "(nombre d'actions, budget total, taux de réalisation).", "MUST"),
        ("F-A1-02", "Filtres sur la liste des actions : commune, type_action, statut, "
                    "priorité, année.", "MUST"),
        ("F-A1-03", "Fiche détaillée d'un plan avec toutes ses actions tabulées.", "MUST"),
        ("F-A1-04", "Tableau de synthèse budgétaire agrégé par commune et par type d'action "
                    "(matrice croisée).", "MUST"),
        ("F-A1-05", "Export Excel du plan annuel complet (.xlsx via openpyxl).", "MUST"),
        ("F-A1-06", "Import depuis un fichier Excel (template téléchargeable).", "SHOULD"),
        ("F-A1-07", "Indicateur de progression globale : barre de pourcentage "
                    "(actions réalisées / total) par plan.", "MUST"),
        ("F-A1-08", "Graphique camembert répartition du budget par type d'action.", "SHOULD"),
    ],
    [1.8, 12.0, 2.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  5. AXE 2 — CALENDRIER D'INTERVENTION
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "5. Axe 2 — Calendrier d'intervention")

# ─── 5.1 Modèle de données ─────────────────────────────────────────────────

heading2(doc, "5.1 Modèle de données")

heading3(doc, "CalendrierIntervention — en-tête du calendrier d'une action")
make_table(doc,
    ["Champ", "Type Django", "Contraintes / Choix"],
    [
        ("action",            "OneToOneField → ActionPlan",
         "Un calendrier par action. CASCADE delete."),
        ("date_debut_prevue", "DateField", "Date de démarrage planifiée."),
        ("date_fin_prevue",   "DateField", "Date de fin planifiée."),
        ("mode_realisation",  "CharField(choices)",
         "etude_interne_ormva / marche_public / appel_manifestation_interet / regie"),
        ("chef_projet",       "FK → Utilisateur",
         "Responsable global du suivi."),
        ("statut_calendrier", "CharField(choices)",
         "brouillon / valide / cloture. Défaut : brouillon."),
        ("valide_par",        "FK → Utilisateur(null=True)",
         "Renseigné uniquement par l'administrateur lors de la validation."),
        ("date_validation",   "DateTimeField(null=True)", "—"),
        ("observations",      "TextField(blank=True)", "—"),
    ],
    [4.0, 4.5, 10.0]
)

heading3(doc, "TacheIntervention — tâche élémentaire")
make_table(doc,
    ["Champ", "Type Django", "Contraintes / Choix"],
    [
        ("calendrier",         "FK → CalendrierIntervention", "CASCADE delete."),
        ("code_tache",         "CharField(20)",
         "Identifiant alphanumérique : T01, T02… Unique par calendrier."),
        ("nom_tache",          "CharField(200)", "Intitulé court de la tâche."),
        ("description",        "TextField(blank=True)", "Détail des travaux ou prestations."),
        ("date_debut_prevue",  "DateField", "—"),
        ("date_fin_prevue",    "DateField", "—"),
        ("duree_prevue",       "PositiveIntegerField",
         "Durée en jours calendaires."),
        ("taches_anterieures", "ManyToManyField('self', blank=True)",
         "Dépendances finish-to-start entre tâches du même calendrier."),
        ("responsable",        "FK → Utilisateur",
         "Utilisateur chargé du suivi de cette tâche (doit exister dans la plateforme)."),
        ("type_suivi",         "CharField(choices)",
         "suivi_travaux / suivi_etude / suivi_administratif / realisation_etude_interne"),
        ("statut_tache",       "CharField(choices)",
         "non_demarree / en_cours / terminee / bloquee. Défaut : non_demarree."),
        ("observations",       "TextField(blank=True)", "—"),
    ],
    [4.0, 4.5, 10.0]
)

note(doc, "Le champ taches_anterieures implémente les dépendances Finish-to-Start (FS) : "
          "une tâche ne peut démarrer qu'une fois TOUTES ses tâches antérieures terminées. "
          "C'est la relation standard utilisée par les diagrammes de Gantt et le calcul PERT/CPM.")

# ─── 5.2 Formulaire et Gantt ───────────────────────────────────────────────

heading2(doc, "5.2 Fonctionnalités — Formulaire et diagramme de Gantt")

make_table(doc,
    ["ID", "Fonctionnalité", "Priorité"],
    [
        ("F-A2-01", "Formulaire de saisie du calendrier : champs de l'en-tête + section "
                    "dynamique d'ajout/suppression de tâches (JavaScript — formset Django).", "MUST"),
        ("F-A2-02", "Dans le formulaire de tâche : champ taches_anterieures affiché sous forme "
                    "de liste à cocher des tâches existantes du même calendrier.", "MUST"),
        ("F-A2-03", "Validation côté serveur : détection des cycles dans le graphe de dépendances "
                    "(algorithme de tri topologique — erreur explicite si cycle détecté).", "MUST"),
        ("F-A2-04", "Vue diagramme de Gantt (Frappe Gantt) :\n"
                    "  • Barre par tâche avec couleur selon statut_tache\n"
                    "  • Liens de dépendance entre tâches (flèches FS)\n"
                    "  • Survol : tooltip avec code, responsable, durée\n"
                    "  • Mise en évidence des tâches du chemin critique (barre rouge)\n"
                    "  • Mode lecture seule pour visiteur et éditeur.", "MUST"),
        ("F-A2-05", "Barre de progression sous le Gantt : date courante marquée par une ligne "
                    "verticale rouge (« aujourd'hui »).", "MUST"),
        ("F-A2-06", "Bouton « Valider le calendrier » visible uniquement par l'administrateur. "
                    "Après validation : aucune modification possible par l'opérateur.", "MUST"),
        ("F-A2-07", "Notification interne à chaque responsable lors de son affectation à une tâche "
                    "(badge dans la barre de navigation — pas d'e-mail en V1).", "SHOULD"),
        ("F-A2-08", "Export PDF du Gantt (capture côté client via html2canvas ou similaire).", "SHOULD"),
    ],
    [1.8, 12.5, 2.0]
)

# ─── 5.3 PERT ──────────────────────────────────────────────────────────────

heading2(doc, "5.3 Vue réseau PERT")

body(doc, (
    "La vue PERT est accessible depuis le calendrier validé. Elle représente les tâches comme "
    "des nœuds d'un graphe orienté (réseau). Elle est implémentée avec vis.js Network et "
    "calculée côté serveur par Django. La vue est en lecture seule."
))

make_table(doc,
    ["ID", "Fonctionnalité", "Priorité"],
    [
        ("F-A2-09", "Affichage du réseau de tâches sous forme de graphe orienté (DAG) : "
                    "chaque nœud représente une tâche, chaque arc représente une dépendance FS.", "MUST"),
        ("F-A2-10", "Chaque nœud affiche : code_tache, nom_tache (court), durée prévue, "
                    "statut (couleur de fond).", "MUST"),
        ("F-A2-11", "Calcul du chemin critique (CPM — Critical Path Method) côté serveur Django :\n"
                    "  1. Passage avant (forward pass) : dates au plus tôt (ES, EF)\n"
                    "  2. Passage arrière (backward pass) : dates au plus tard (LS, LF)\n"
                    "  3. Marge totale = LS − ES\n"
                    "  4. Chemin critique = tâches avec marge totale = 0.", "MUST"),
        ("F-A2-12", "Mise en évidence du chemin critique : arcs et nœuds en rouge vif. "
                    "Durée totale du projet affichée en haut de la vue.", "MUST"),
        ("F-A2-13", "Tooltip au survol d'un nœud : ES, EF, LS, LF, marge totale.", "SHOULD"),
        ("F-A2-14", "Disposition automatique du graphe (layout hiérarchique de gauche à droite).", "MUST"),
    ],
    [1.8, 12.5, 2.0]
)

note(doc, "Le calcul CPM est exposé par un endpoint Django /plan/action/<id>/pert/data/ "
          "qui retourne les nœuds et arcs enrichis des dates ES/EF/LS/LF. "
          "vis.js consomme ce JSON pour le rendu graphique.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  6. AXE 3 — SUIVI D'AVANCEMENT
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "6. Axe 3 — Suivi d'avancement")

# ─── 6.1 Modèle de données ─────────────────────────────────────────────────

heading2(doc, "6.1 Modèle de données")

heading3(doc, "SuiviAvancement — rapport périodique par tâche")
make_table(doc,
    ["Champ", "Type Django", "Contraintes / Choix"],
    [
        ("tache",                  "FK → TacheIntervention", "CASCADE delete."),
        ("auteur",                 "FK → Utilisateur",
         "Doit être égal à tache.responsable (contrôle côté serveur)."),
        ("date_rapport",           "DateField", "Date du rapport d'avancement."),
        ("avancement_pct",         "PositiveIntegerField",
         "0 à 100. Pourcentage d'avancement physique de la tâche."),
        ("etat_bloc",              "CharField(choices)",
         "conforme / retard / bloque / termine"),
        ("commentaire",            "TextField(blank=True)",
         "Description de l'avancement, difficultés rencontrées."),
        ("date_prochaine_echeance","DateField(null=True, blank=True)",
         "Prochaine étape planifiée."),
        ("date_saisie",            "DateTimeField(auto_now_add=True)", "—"),
    ],
    [4.5, 4.5, 9.5]
)

heading3(doc, "PieceJustificative — document rattaché à un rapport")
make_table(doc,
    ["Champ", "Type Django", "Contraintes / Choix"],
    [
        ("suivi",        "FK → SuiviAvancement", "CASCADE delete."),
        ("type_piece",   "CharField(choices)",
         "pv_attachement / pv_reception / photo_chantier / rapport_etude / "
         "note_administrative / autre"),
        ("fichier",      "FileField",
         "Stocké dans MEDIA_ROOT/plan_action/pieces/<annee>/<action_id>/. "
         "Formats : PDF, JPG, PNG, DOCX. Taille max : 10 Mo."),
        ("libelle",      "CharField(200)",
         "Titre court du document (ex. : « PV attachement phase 1 »)."),
        ("date_document","DateField(null=True, blank=True)",
         "Date du document (utile pour PV et rapports)."),
        ("date_upload",  "DateTimeField(auto_now_add=True)", "—"),
        ("uploade_par",  "FK → Utilisateur", "Enregistré automatiquement."),
    ],
    [3.5, 4.0, 11.0]
)

# ─── 6.2 Fonctionnalités ───────────────────────────────────────────────────

heading2(doc, "6.2 Fonctionnalités")

make_table(doc,
    ["ID", "Fonctionnalité", "Priorité"],
    [
        ("F-A3-01", "Tableau de bord du suivi par action : liste des tâches avec barre de "
                    "progression (avancement_pct), badge etat_bloc coloré, responsable, "
                    "date du dernier rapport.", "MUST"),
        ("F-A3-02", "Formulaire de saisie d'un rapport d'avancement, accessible uniquement "
                    "à l'utilisateur identifié comme responsable de la tâche dans le "
                    "calendrier validé.", "MUST"),
        ("F-A3-03", "Upload multi-fichiers de pièces justificatives (glisser-déposer "
                    "via Dropzone.js) : contrôle format et taille avant envoi.", "MUST"),
        ("F-A3-04", "Galerie photos : miniatures des photos uploadées avec visionneuse "
                    "plein écran (Lightbox2). Filtrée par type_piece = photo_chantier.", "MUST"),
        ("F-A3-05", "Courbe S (S-curve) : graphique Chart.js superposant l'avancement "
                    "prévu (linéaire théorique) et l'avancement réel cumulé par date "
                    "de rapport.", "SHOULD"),
        ("F-A3-06", "Timeline historique : liste chronologique de tous les rapports d'une "
                    "tâche avec résumé (date, %, état, commentaire) et liens vers les "
                    "pièces jointes.", "MUST"),
        ("F-A3-07", "Alerte visuelle retard : badge rouge sur toute tâche dont "
                    "l'avancement réel est inférieur à l'avancement théorique attendu "
                    "à la date courante (calculé par interpolation linéaire).", "MUST"),
        ("F-A3-08", "Vue synthèse globale (niveau plan) : tableau récapitulatif de toutes "
                    "les actions du plan avec taux d'avancement moyen pondéré.", "SHOULD"),
    ],
    [1.8, 12.5, 2.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  7. MATRICE DES DROITS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "7. Matrice des droits")

body(doc, (
    "La matrice ci-dessous précise les permissions par rôle pour chaque axe fonctionnel. "
    "Elle s'appuie sur le modèle compte.Utilisateur avec les rôles : "
    "visiteur / opérateur / éditeur, et étend les droits à l'administrateur Django (superuser)."
))

make_rights_table(doc,
    ["Action", "Visiteur", "Opérateur", "Éditeur", "Administrateur"],
    [
        # Axe 1
        ("A1 — Consulter les plans et actions", "✓", "✓", "✓", "✓"),
        ("A1 — Créer / modifier un plan",        "✗", "✓", "✗", "✓"),
        ("A1 — Supprimer un plan",               "✗", "✗", "✗", "✓"),
        # Axe 2
        ("A2 — Consulter le calendrier et le Gantt", "✓", "✓", "✓", "✓"),
        ("A2 — Créer un calendrier (brouillon)",  "✗", "✓", "✗", "✓"),
        ("A2 — Modifier les tâches (brouillon)",  "✗", "✓", "✗", "✓"),
        ("A2 — Valider le calendrier",            "✗", "✗", "✗", "✓"),
        ("A2 — Modifier un calendrier validé",    "✗", "✗", "✗", "✓"),
        ("A2 — Consulter la vue PERT",            "✓", "✓", "✓", "✓"),
        # Axe 3
        ("A3 — Consulter le suivi d'avancement",  "✓", "✓", "✓", "✓"),
        ("A3 — Saisir rapport (propre tâche) (*)", "✗", "✓", "✓", "✓"),
        ("A3 — Upload pièces justificatives (*)",  "✗", "✓", "✓", "✓"),
        ("A3 — Modifier rapport d'un autre",      "✗", "✗", "✗", "✓"),
        ("A3 — Supprimer une pièce jointe",       "✗", "✗", "✗", "✓"),
    ],
    [7.0, 2.0, 2.3, 2.0, 3.5]
)

note(doc, "(*) La saisie d'un rapport et l'upload de pièces justificatives sont autorisés "
          "à l'opérateur et à l'éditeur UNIQUEMENT si l'utilisateur est le responsable "
          "de la tâche concernée dans le calendrier validé. Ce contrôle est effectué "
          "côté serveur dans la vue Django, pas uniquement en template.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  8. ARCHITECTURE TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "8. Architecture technique")

heading2(doc, "8.1 Structure de la nouvelle app Django")

body(doc, "Nouvelle application : plan_action/ — à créer dans le répertoire plateformeSIG/")

make_table(doc,
    ["Fichier", "Contenu"],
    [
        ("models.py",
         "PlanAmenagement, ActionPlan, CalendrierIntervention, "
         "TacheIntervention (M2M self), SuiviAvancement, PieceJustificative"),
        ("views.py",
         "Vues CRUD des 3 axes + vue Gantt data (JSON) + vue PERT data (JSON) + "
         "vue validation calendrier + vue upload pièces + vue export Excel"),
        ("urls.py",         "Préfixe : /plan/"),
        ("forms.py",        "PlanForm, ActionPlanForm, CalendrierForm + TacheFormSet (extra=1), "
                            "SuiviAvancementForm, PieceJustificativeForm"),
        ("admin.py",        "Enregistrement de tous les modèles avec actions inline"),
        ("utils.py",        "Algorithme CPM (forward/backward pass), calcul retard, calcul taux"),
        ("tests.py",        "Tests unitaires modèles + tests vues permissions"),
        ("templates/plan_action/", "Dossier des templates (détail ci-dessous)"),
        ("static/plan_action/",    "JavaScript spécifique Gantt, PERT, upload"),
    ],
    [5.0, 13.5]
)

heading2(doc, "8.2 Templates")
make_table(doc,
    ["Template", "Description"],
    [
        ("plan_list.html",          "Liste des plans annuels + indicateurs synthétiques"),
        ("plan_detail.html",        "Fiche plan : tableau des actions avec filtres"),
        ("plan_form.html",          "Formulaire création / modification plan"),
        ("action_form.html",        "Formulaire action (création dans un plan)"),
        ("calendrier_form.html",    "Formulaire calendrier + formset tâches dynamique"),
        ("calendrier_gantt.html",   "Vue Gantt (Frappe Gantt) + barre aujourd'hui"),
        ("calendrier_pert.html",    "Vue réseau PERT (vis.js) + chemin critique"),
        ("suivi_dashboard.html",    "Dashboard suivi : liste tâches + barres avancement"),
        ("suivi_form.html",         "Formulaire rapport d'avancement + upload pièces"),
        ("suivi_historique.html",   "Timeline chronologique des rapports d'une tâche"),
    ],
    [6.0, 12.5]
)

heading2(doc, "8.3 URLs principales")
make_table(doc,
    ["URL", "Vue", "Description"],
    [
        ("/plan/",                         "plan_list",          "Liste des plans annuels"),
        ("/plan/creer/",                   "plan_create",        "Formulaire création plan"),
        ("/plan/<id>/",                    "plan_detail",        "Détail + liste actions du plan"),
        ("/plan/<id>/modifier/",           "plan_update",        "Modifier un plan"),
        ("/plan/<id>/action/ajouter/",     "action_create",      "Ajouter une action au plan"),
        ("/plan/action/<id>/calendrier/",  "calendrier_form",    "Créer / éditer calendrier"),
        ("/plan/action/<id>/gantt/",       "calendrier_gantt",   "Afficher le Gantt"),
        ("/plan/action/<id>/gantt/data/",  "gantt_data",         "Endpoint JSON pour Frappe Gantt"),
        ("/plan/action/<id>/pert/",        "calendrier_pert",    "Afficher le réseau PERT"),
        ("/plan/action/<id>/pert/data/",   "pert_data",          "Endpoint JSON CPM pour vis.js"),
        ("/plan/action/<id>/valider/",     "valider_calendrier", "Validation admin (POST)"),
        ("/plan/tache/<id>/suivi/",        "suivi_form",         "Saisir un rapport"),
        ("/plan/tache/<id>/suivi/liste/",  "suivi_historique",   "Timeline rapports"),
        ("/plan/<id>/export/excel/",       "export_plan_excel",  "Export Excel plan"),
    ],
    [5.5, 4.0, 9.0]
)

heading2(doc, "8.4 Algorithme CPM (utils.py)")
body(doc, (
    "L'algorithme CPM est implémenté en Python pur dans plan_action/utils.py. "
    "Il reçoit la liste des TacheIntervention d'un calendrier et retourne pour chaque tâche : "
    "ES (Early Start), EF (Early Finish), LS (Late Start), LF (Late Finish), marge totale, "
    "et un booléen is_critical."
))
bullet(doc, "Étape 1 — Tri topologique (Kahn's algorithm) pour ordonner les tâches "
            "sans violer les dépendances.")
bullet(doc, "Étape 2 — Forward pass : calcul de ES = max(EF des prédécesseurs) et EF = ES + durée.")
bullet(doc, "Étape 3 — Backward pass : calcul de LF = min(LS des successeurs) et LS = LF − durée.")
bullet(doc, "Étape 4 — Marge totale = LS − ES. Chemin critique = tâches avec marge = 0.")
note(doc, "Si le graphe contient un cycle (erreur de saisie), une ValueError est levée "
          "avec le message d'erreur affiché dans le formulaire.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  9. EXIGENCES NON FONCTIONNELLES
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "9. Exigences non fonctionnelles")

heading2(doc, "9.1 Performance")
make_table(doc,
    ["ID", "Exigence", "Cible"],
    [
        ("PERF-01", "Chargement de la liste d'un plan avec 50 actions et filtres", "< 1 s"),
        ("PERF-02", "Rendu du Gantt avec 20 tâches (données JSON + rendu JS)", "< 2 s"),
        ("PERF-03", "Calcul CPM côté serveur (20 tâches, graphe dense)", "< 0.2 s"),
        ("PERF-04", "Upload d'un fichier de 10 Mo", "< 5 s"),
        ("PERF-05", "Génération export Excel d'un plan avec 100 actions", "< 3 s"),
    ],
    [1.5, 12.0, 2.5]
)

heading2(doc, "9.2 Sécurité")
make_table(doc,
    ["ID", "Exigence"],
    [
        ("SEC-01", "Tous les endpoints /plan/* doivent exiger @login_required."),
        ("SEC-02", "Le contrôle « responsable de la tâche » doit être vérifié côté serveur "
                   "(get_object_or_404 + assertion sur request.user == tache.responsable). "
                   "Le template masquant le bouton n'est PAS suffisant."),
        ("SEC-03", "Validation du type MIME des fichiers uploadés côté serveur "
                   "(whitelist : application/pdf, image/jpeg, image/png, "
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document)."),
        ("SEC-04", "Les fichiers médias sont servis par Django en développement "
                   "(MEDIA_URL) — configurer un accès contrôlé en production "
                   "(nginx + token ou X-Accel-Redirect)."),
        ("SEC-05", "Protection CSRF sur tous les formulaires POST (middleware Django actif)."),
        ("SEC-06", "Toute modification d'un CalendrierIntervention validé est bloquée "
                   "sauf pour le superuser (vérification dans la vue update)."),
    ],
    [1.5, 17.0]
)

heading2(doc, "9.3 Stockage des fichiers")
bullet(doc, "Chemin physique : MEDIA_ROOT/plan_action/pieces/<annee>/<action_id>/<nom_fichier>")
bullet(doc, "Taille maximale par fichier : 10 Mo (contrôle via Django FileField + validators).")
bullet(doc, "Formats autorisés : PDF, JPG, PNG, DOCX.")
bullet(doc, "Les fichiers supprimés par l'administrateur sont physiquement effacés "
            "(signal post_delete sur PieceJustificative).")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  10. CRITÈRES D'ACCEPTATION
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "10. Critères d'acceptation")

make_table(doc,
    ["ID", "Scénario de test", "Résultat attendu"],
    [
        ("SC-01",
         "L'administrateur crée un plan 2026 avec 5 actions dans 3 communes différentes.",
         "Plan visible dans la liste ; filtrable par commune et par type d'action ; "
         "indicateur budget et taux à 0 %."),
        ("SC-02",
         "L'opérateur crée un calendrier pour l'action ACT-01, avec 4 tâches (T1, T2, T3, T4) "
         "où T3 dépend de T1 et T2, et T4 dépend de T3.",
         "Liens de dépendance visibles sur le Gantt ; ordre de tri topologique respecté."),
        ("SC-03",
         "L'opérateur tente de créer une dépendance cyclique (T2 dépend de T3, T3 dépend de T2).",
         "Erreur explicite affichée dans le formulaire : « Dépendance cyclique détectée. »"),
        ("SC-04",
         "L'administrateur valide le calendrier.",
         "Statut passe à « validé » ; le formulaire d'édition des tâches est désactivé "
         "pour l'opérateur."),
        ("SC-05",
         "Vue PERT calculée pour le calendrier de SC-02.",
         "Chemin critique T1→T3→T4 ou T2→T3→T4 en rouge ; durée totale affichée ; "
         "tooltip sur T3 affiche ES, EF, LS, LF, marge = 0."),
        ("SC-06",
         "Le responsable de T2 saisit un rapport à 40 % avec une photo chantier.",
         "Rapport enregistré ; barre d'avancement de T2 = 40 % sur le dashboard ; "
         "photo visible dans la galerie."),
        ("SC-07",
         "Un opérateur non-responsable tente d'accéder au formulaire de rapport de T1.",
         "HTTP 403 ou redirection avec message d'erreur."),
        ("SC-08",
         "Tâche T2 est en retard (avancement réel < avancement théorique à date courante).",
         "Badge rouge visible sur le dashboard suivi et sur le Gantt."),
        ("SC-09",
         "Visiteur tente d'accéder à /plan/creer/ et à /plan/tache/1/suivi/.",
         "Redirection vers /compte/login/ dans les deux cas."),
        ("SC-10",
         "L'administrateur exporte le plan 2026 en Excel.",
         "Fichier .xlsx téléchargé contenant toutes les actions avec commune, "
         "type, budget, statut, priorité."),
    ],
    [1.5, 7.5, 9.5]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  11. LIVRABLES ET PLANNING INDICATIF
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "11. Livrables et planning indicatif")

make_table(doc,
    ["Semaine", "Livrable", "Axe"],
    [
        ("S1",     "Modèles Django + migrations initiales + enregistrement admin", "A1/A2/A3"),
        ("S2",     "CRUD Axe 1 : création/modification/suppression plans et actions + filtres", "A1"),
        ("S3",     "Tableau de synthèse budgétaire + indicateurs + export Excel plan", "A1"),
        ("S4",     "CRUD Axe 2 : formulaire calendrier + formset tâches dynamique + "
                   "détection cycles", "A2"),
        ("S5",     "Diagramme de Gantt (Frappe Gantt) + données JSON + barre aujourd'hui", "A2"),
        ("S6",     "Vue réseau PERT (vis.js) + algorithme CPM + mise en évidence "
                   "chemin critique", "A2"),
        ("S7",     "CRUD Axe 3 : formulaire suivi + upload multi-fichiers + "
                   "galerie photos", "A3"),
        ("S8",     "Courbe S (Chart.js) + timeline historique + alertes retard", "A3"),
        ("S9",     "Matrice des droits complète + tests permissions + sécurité fichiers", "A1/A2/A3"),
        ("S10",    "Tests d'acceptation (10 scénarios SC-01 à SC-10) + corrections + "
                   "documentation utilisateur", "A1/A2/A3"),
    ],
    [1.8, 13.5, 3.0]
)

note(doc, "Planning indicatif sur 10 semaines — à ajuster selon les priorités du PFE. "
          "La V1 prioritaire couvre les fonctionnalités MUST. "
          "Les fonctionnalités SHOULD (courbe S, export PDF Gantt, notification interne) "
          "sont planifiées en S7–S9 selon l'avancement.")


# ══════════════════════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"Document généré : {OUTPUT}")
