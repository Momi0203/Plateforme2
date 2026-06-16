"""
Génération du Cahier des Charges — Module Carte de la Plateforme SIG
Exécuter depuis n'importe où : python generate_cdc.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "CDC_Module_Carte_PlateformeSIG_v2.docx")

# ── Couleurs ─────────────────────────────────────────────────────────────────
C_BLEU_TITRE  = RGBColor(0x1F, 0x49, 0x7D)   # bleu DGA
C_BLEU_H2     = RGBColor(0x2E, 0x74, 0xB5)
C_BLEU_H3     = RGBColor(0x2F, 0x5A, 0x8C)
C_ORANGE      = RGBColor(0xC5, 0x5A, 0x11)
C_VERT        = RGBColor(0x37, 0x86, 0x30)
C_GRIS_CELL   = RGBColor(0xBF, 0xD7, 0xED)   # entête table claire
C_BLEU_FONCE  = RGBColor(0x1F, 0x49, 0x7D)
C_BLANC       = RGBColor(0xFF, 0xFF, 0xFF)
C_GRIS_ALT    = RGBColor(0xF2, 0xF2, 0xF2)


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '4472C4')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break()


def heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = C_BLEU_TITRE
    run.font.size = Pt(16)
    run.bold = True
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = C_BLEU_H2
    run.font.size = Pt(13)
    run.bold = True
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = C_BLEU_H3
    run.font.size = Pt(11)
    run.bold = True
    return p


def body(doc, text, indent=0):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("Note : ")
    run.bold = True
    run.font.color.rgb = C_ORANGE
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = C_ORANGE
    p.paragraph_format.space_after = Pt(6)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Entête
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_BLEU_FONCE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_BLANC
        run.font.size = Pt(10)

    # Données
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = C_GRIS_ALT if ri % 2 == 0 else C_BLANC
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)

    # Largeurs colonnes
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = Cm(col_widths[ci])

    set_cell_borders(table)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Police par défaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("PLATEFORME SIG — GESTION DES RESSOURCES EN EAU")
run.font.size = Pt(14)
run.font.color.rgb = C_BLEU_TITRE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Périmètres irrigués — Tafilalet / Midelt")
run.font.size = Pt(11)
run.font.color.rgb = C_BLEU_H2

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CAHIER DES CHARGES")
run.font.size = Pt(28)
run.font.color.rgb = C_BLEU_TITRE
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MODULE « CARTE »")
run.font.size = Pt(22)
run.font.color.rgb = C_ORANGE
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Version",          "1.0"),
    ("Date",             "Juin 2026"),
    ("Statut",           "Projet — en attente de validation"),
    ("Plateforme",       "Django 6 / GeoDjango / PostGIS"),
    ("Environnement",    "Windows 11 — OSGeo4W — PostgreSQL 15"),
    ("Auteur",           "Équipe de développement PFE"),
]
table = doc.add_table(rows=len(meta), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta):
    row = table.rows[i]
    c0 = row.cells[0]
    c1 = row.cells[1]
    set_cell_bg(c0, C_GRIS_CELL)
    set_cell_bg(c1, C_BLANC)
    run_k = c0.paragraphs[0].add_run(k)
    run_k.bold = True
    run_k.font.size = Pt(10)
    run_v = c1.paragraphs[0].add_run(v)
    run_v.font.size = Pt(10)
    c0.width = Cm(5)
    c1.width = Cm(8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SOMMAIRE (manuel)
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "Table des matières")
toc_items = [
    ("1.", "Présentation générale et contexte"),
    ("2.", "Objectifs du module Carte"),
    ("3.", "Périmètre fonctionnel"),
    ("4.", "Inventaire des couches géographiques"),
    ("5.", "Spécifications fonctionnelles détaillées"),
    ("  5.1", "Panneau gauche — Contrôle de visualisation"),
    ("  5.2", "Zone centrale — Trois modes de présentation"),
    ("  5.3", "Panneau droit — Outils et calculs avancés"),
    ("6.", "Interactions et synchronisation entre vues"),
    ("7.", "Interfaces de programmation (API Django)"),
    ("8.", "Sécurité et gestion des droits"),
    ("9.", "Exigences non fonctionnelles"),
    ("10.", "Architecture technique cible"),
    ("11.", "Contraintes et dépendances"),
    ("12.", "Critères d'acceptation"),
    ("13.", "Livrables attendus"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 if num.startswith(" ") else 0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{num.strip()}  {title}")
    run.font.size = Pt(10.5)
    if not num.startswith(" "):
        run.bold = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  1. PRÉSENTATION GÉNÉRALE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "1. Présentation générale et contexte")

heading2(doc, "1.1 Contexte du projet")
body(doc, (
    "La Plateforme SIG est une application web développée avec Django 6 et GeoDjango (PostGIS) "
    "pour la gestion des ressources en eau dans les périmètres irrigués du Sud-Est marocain "
    "(région Tafilalet — provinces de Midelt et Errachidia). Elle centralise quatre domaines "
    "métier distincts :"
))
bullet(doc, "Analyse hydrologique : bassins versants, stations pluviométriques et hydrométriques, "
            "crues, coefficients de Montana.")
bullet(doc, "Diagnostic des ouvrages hydrauliques : seuils, murs de protection, séguias, "
            "barrages de retenue, khettaras, forages/puits et prises locales.")
bullet(doc, "Bilan besoins-ressources en eau : calcul mensuel (calendrier Sep→Aoû) des besoins "
            "agricoles et des ressources disponibles par périmètre.")
bullet(doc, "Géographies de référence : provinces et communes avec données démographiques "
            "et climatiques.")
body(doc, (
    "Le module « Carte » est le module transversal manquant : il doit offrir une vue "
    "cartographique unifiée permettant de visualiser, interroger, analyser et restituer "
    "l'ensemble des entités géographiques stockées dans la base PostGIS."
))

heading2(doc, "1.2 Base de données existante")
body(doc, (
    "La base de données PostgreSQL/PostGIS est opérationnelle. Elle contient les schémas "
    "suivants (correspondant aux cinq applications Django) :"
))
make_table(doc,
    ["Application Django", "Tables principales", "Géométries disponibles"],
    [
        ("carte", "Province, Commune", "PolygonField (SRID 4326)"),
        ("analyse_hydrologique",
         "BassinVersant, StationPluviometrique,\nStationHydrometrique, ReseauHydrographique,\nCoefficientMontana, ResultatAnalyseHydrologique",
         "PolygonField, PolygonField,\nPointField, LineStringField"),
        ("diagnostic",
         "Perimetre, Seuil, MurProtection,\nSeguias/TronconSeguia, BarrageRetenue,\nKhettara, ForagePuits, PriseLocale",
         "GeometryField, PointField,\nGeometryField, GeometryField,\nGeometryField, GeometryField,\nGeometryField, GeometryField"),
        ("Besions_Ressources", "StationClimatique, BilanBesoinRessources", "PointField"),
        ("compte", "Utilisateur (rôles : visiteur / opérateur / éditeur)", "—"),
    ],
    [5.5, 7.5, 5.5]
)
note(doc, "Toutes les géométries vectorielles sont stockées en SRID 4326 (WGS 84). "
          "Les coordonnées X/Y dans les champs textuels sont en projection Nord Maroc (EPSG:26191).")

heading2(doc, "1.3 Pile technologique")
make_table(doc,
    ["Composant", "Technologie / Version"],
    [
        ("Framework web",       "Django 6.0.4"),
        ("Base de données",     "PostgreSQL 15 + PostGIS (extension GIS)"),
        ("Bibliothèque GIS",    "GeoDjango (django.contrib.gis)"),
        ("Rendu cartographique front-end", "MapLibre GL JS (ou Leaflet selon décision d'implémentation)"),
        ("Bibliothèques GIS natives", "GDAL 3.12 / GEOS / PROJ — OSGeo4W"),
        ("Serveur de tuiles (option)", "GeoServer ou pg_tileserv"),
        ("Export de graphiques", "matplotlib / Chart.js"),
        ("Export Excel",        "openpyxl"),
        ("Export PDF/carte",    "WeasyPrint ou reportlab"),
        ("Gestion utilisateurs", "compte.Utilisateur (AbstractUser + rôle)"),
    ],
    [7, 11]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  2. OBJECTIFS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "2. Objectifs du module Carte")

heading2(doc, "2.1 Objectif général")
body(doc, (
    "Développer un module de visualisation et d'analyse géospatiale intégré à la plateforme "
    "existante, permettant à l'utilisateur de naviguer sur une carte interactive, d'interroger "
    "les entités géographiques de la base de données, d'effectuer des analyses spatiales simples "
    "et avancées, et de restituer les résultats sous trois formes complémentaires : "
    "carte, tableau de bord synthétique et tableau attributaire."
))

heading2(doc, "2.2 Objectifs spécifiques")
bullet(doc, "OBJ-1 : Afficher en superposition toutes les couches géographiques de la base PostGIS "
            "(provinces, communes, bassins versants, périmètres, ouvrages hydrauliques, réseaux "
            "hydrographiques, stations).")
bullet(doc, "OBJ-2 : Permettre le contrôle fin de la visualisation (visibilité, ordre, symbologie "
            "personnalisée) depuis un panneau gauche dédié.")
bullet(doc, "OBJ-3 : Fournir des outils de sélection et de requête attributaire et spatiale, "
            "simples et multicritères.")
bullet(doc, "OBJ-4 : Mettre à disposition des outils d'analyse géospatiale (tampon, intersection, "
            "proximité, statistiques de zone) dans un panneau droit inspiré des boîtes à outils SIG.")
bullet(doc, "OBJ-5 : Proposer trois modes de restitution synchronisés — carte exportable, dashboard "
            "de synthèse et tableau attributaire — depuis la même zone de travail centrale.")
bullet(doc, "OBJ-6 : Permettre l'export des résultats en formats imprimables (PDF A4–A0) et "
            "numériques (PNG, CSV, Excel, GeoJSON).")
bullet(doc, "OBJ-7 : Respecter la matrice des rôles existante (visiteur / opérateur / éditeur) "
            "pour l'accès aux fonctionnalités du module.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  3. PÉRIMÈTRE FONCTIONNEL
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "3. Périmètre fonctionnel")

heading2(doc, "3.1 Dans le périmètre")
make_table(doc,
    ["ID", "Fonctionnalité", "Priorité"],
    [
        ("F-01", "Affichage cartographique des couches PostGIS", "MUST"),
        ("F-02", "Gestionnaire de couches (arborescence, visibilité, ordre)", "MUST"),
        ("F-03", "Sélection d'entités (rectangle, polygone, clic)", "MUST"),
        ("F-04", "Requête attributaire simple (une condition)", "MUST"),
        ("F-05", "Requête multicritère avec opérateurs logiques AND/OR", "MUST"),
        ("F-06", "Symbologie simple (couleur, taille, opacité)", "MUST"),
        ("F-07", "Symbologie catégorisée et graduée", "MUST"),
        ("F-08", "Bibliothèque de styles personnels (sauvegarde/réutilisation)", "SHOULD"),
        ("F-09", "Outils d'analyse : tampon, intersection, proximité", "MUST"),
        ("F-10", "Outils de gestion : fusion (merge), dissolution, jointure spatiale", "SHOULD"),
        ("F-11", "Calculatrice de champ (expressions sur les attributs)", "SHOULD"),
        ("F-12", "Requête SQL avancée avec fonctions spatiales PostGIS", "COULD"),
        ("F-13", "Vue carte avec outil d'exportation (A4–A0, PDF/PNG)", "MUST"),
        ("F-14", "Vue dashboard avec graphiques liés à la sélection", "MUST"),
        ("F-15", "Vue tableau attributaire avec tri/filtre/pagination", "MUST"),
        ("F-16", "Export tableau en CSV / Excel / JSON", "MUST"),
        ("F-17", "Export dashboard en PDF / PNG", "SHOULD"),
        ("F-18", "Synchronisation bidirectionnelle carte ↔ tableau ↔ dashboard", "MUST"),
        ("F-19", "Fond de carte interchangeable (OSM, satellite, sans fond)", "SHOULD"),
        ("F-20", "Import de couche shapefile additionnel (ZIP)", "COULD"),
    ],
    [1.5, 10.5, 2]
)
note(doc, "MUST = requis pour la V1 ; SHOULD = à inclure si le calendrier le permet ; "
          "COULD = version ultérieure.")

heading2(doc, "3.2 Hors périmètre (V1)")
bullet(doc, "Édition géométrique (dessin, modification de formes sur la carte).")
bullet(doc, "Module Efficiences (spécifié dans ETAT_AVANCEMENT_EFFICIENCES.md — développement séparé).")
bullet(doc, "Synchronisation temps réel multi-utilisateur.")
bullet(doc, "Publication de services WMS/WFS externes.")
bullet(doc, "Application mobile.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  4. INVENTAIRE DES COUCHES
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "4. Inventaire des couches géographiques")

body(doc, (
    "Le tableau ci-dessous liste toutes les couches qui devront être exposées via l'API "
    "GeoJSON du module Carte. Chaque couche correspond à un ou plusieurs modèles Django "
    "possédant un champ géométrique dans la base PostGIS."
))

make_table(doc,
    ["Couche (nom affiché)", "Modèle Django", "Type géom.", "Champs attributaires clés", "Groupe"],
    [
        ("Provinces",
         "carte.Province",
         "Polygone",
         "nom_fr, nom_ar, superficie_km2, population_totale, temp_moy_annuelle_c, precip_annuelle_mm, et0_annuelle_mm",
         "Administratif"),
        ("Communes",
         "carte.Commune",
         "Polygone",
         "nom_fr, type_commune, population_totale, superficie_km2, nbr_perimetres_agricoles",
         "Administratif"),
        ("Bassins versants",
         "analyse_hydrologique.BassinVersant",
         "Polygone",
         "nom, surface (km²), perimetre (km), z_min, z_max, thalweg, ouvrage_en_tete",
         "Hydrologie"),
        ("Réseau hydrographique",
         "analyse_hydrologique.ReseauHydrographique",
         "Polyligne",
         "grid_code",
         "Hydrologie"),
        ("Stations pluviométriques",
         "analyse_hydrologique.StationPluviometrique",
         "Polygone (Thiessen)",
         "nom, hauteur_moyenne (mm), pjmax_t10/20/50/100",
         "Hydrologie"),
        ("Stations hydrométriques",
         "analyse_hydrologique.StationHydrometrique",
         "Point",
         "nom, superficie_bv_jaugee, qjmax_t10/20/50/100",
         "Hydrologie"),
        ("Stations climatiques",
         "Besions_Ressources.StationClimatique",
         "Point",
         "nom, latitude, temperatures_moyennes (JSON), precipitations_normales (JSON)",
         "Hydrologie"),
        ("Périmètres agricoles",
         "diagnostic.Perimetre",
         "Géométrie",
         "ksar_village, commune, superficie_totale, superficie_irriguee, nombre_beneficiaires, statut",
         "Diagnostic"),
        ("Seuils",
         "diagnostic.Seuil",
         "Point",
         "nom_du_seuil, nature, type, debit_mobilise (l/s), longueur, hauteur, etat_construction, statut",
         "Diagnostic"),
        ("Murs de protection",
         "diagnostic.MurProtection",
         "Géométrie",
         "nom_mur_protection, rive, position, nature_materiaux, longueur, hauteur, statut",
         "Diagnostic"),
        ("Tronçons de séguias",
         "diagnostic.TronconSeguia",
         "Géométrie",
         "seguia.nom, troncon, longueur, nature, debit (m³/s), efficience_calculee, statut",
         "Diagnostic"),
        ("Barrages de retenue",
         "diagnostic.BarrageRetenue",
         "Géométrie",
         "nom, capacite_retenue, debit_derive, longueur, hauteur, statut",
         "Diagnostic"),
        ("Khettaras",
         "diagnostic.Khettara",
         "Géométrie",
         "nom, debit, longueur, largeur, materiaux_de_construction, statut",
         "Diagnostic"),
        ("Forages / Puits",
         "diagnostic.ForagePuits",
         "Géométrie",
         "nom, debit (m³/h), profondeur, diametre, source_energie_pompage, statut",
         "Diagnostic"),
        ("Prises locales",
         "diagnostic.PriseLocale",
         "Géométrie",
         "nom, forme_pertuis, debit_derive (m³/s), materiaux_construction, statut",
         "Diagnostic"),
    ],
    [3.5, 4.5, 2.0, 6.5, 2.5]
)

note(doc, "Les couches du groupe « Diagnostic » disposent en plus d'un sous-modèle Etat<X> "
          "contenant les notes de diagnostic (0–5) et l'état général. Ces données doivent "
          "être joignables pour la symbologie et le dashboard.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  5. SPÉCIFICATIONS FONCTIONNELLES
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "5. Spécifications fonctionnelles détaillées")

# ─── 5.1 Panneau gauche ────────────────────────────────────────────────────

heading2(doc, "5.1 Panneau gauche — Contrôle de visualisation")

body(doc, (
    "Le panneau gauche est un panneau vertical redimensionnable (largeur par défaut : 280 px, "
    "réductible à une icône). Il contient quatre onglets principaux."
))

heading3(doc, "5.1.1 Onglet « Couches »")
body(doc, "Exigences fonctionnelles :")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("PG-01", "Afficher un arbre hiérarchique des couches organisé par groupes : "
                  "Administratif, Hydrologie, Diagnostic.", "MUST"),
        ("PG-02", "Chaque couche dispose d'une case à cocher pour basculer sa visibilité "
                  "sans supprimer son rendu.", "MUST"),
        ("PG-03", "Glisser-déposer pour réordonner les couches et modifier l'ordre de rendu.", "SHOULD"),
        ("PG-04", "Icône indicatrice du type géométrique (point, ligne, polygone, raster) "
                  "à gauche du nom de couche.", "MUST"),
        ("PG-05", "Badge rouge sur la couche si l'API renvoie une erreur de chargement.", "SHOULD"),
        ("PG-06", "Menu contextuel sur clic droit : Zoomer vers la couche, "
                  "Voir la table attributaire, Ouvrir la symbologie.", "MUST"),
        ("PG-07", "Indicateur du nombre d'entités chargées par couche (tooltip).", "COULD"),
    ],
    [1.8, 11.0, 2.0]
)

heading3(doc, "5.1.2 Onglet « Sélection »")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("SEL-01", "Outil de sélection rectangulaire (drag) sur la carte.", "MUST"),
        ("SEL-02", "Outil de sélection circulaire (centre + rayon).", "SHOULD"),
        ("SEL-03", "Outil de sélection polygonale libre (clic point par point).", "SHOULD"),
        ("SEL-04", "Sélection par clic simple avec info-bulle affichant les attributs "
                   "principaux de l'entité.", "MUST"),
        ("SEL-05", "Modes de sélection : Nouvelle / Ajouter à / Soustraire de / Intersecter avec.", "MUST"),
        ("SEL-06", "Compteur « N entité(s) sélectionnée(s) » mis à jour en temps réel "
                   "dans la barre de statut.", "MUST"),
        ("SEL-07", "Boutons : Tout sélectionner / Tout désélectionner / Inverser la sélection.", "MUST"),
        ("SEL-08", "Mise en surbrillance (highlight) des entités sélectionnées sur la carte "
                   "(couleur jaune/orange distincte).", "MUST"),
    ],
    [1.8, 11.0, 2.0]
)

heading3(doc, "5.1.3 Onglet « Requête simple »")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("RS-01", "Formulaire guidé : sélectionner une couche → choisir un champ "
                  "→ choisir un opérateur → saisir une valeur.", "MUST"),
        ("RS-02", "Opérateurs disponibles : = , ≠ , > , ≥ , < , ≤ , CONTIENT , COMMENCE PAR , "
                  "EST NULL , ENTRE [val1] ET [val2].", "MUST"),
        ("RS-03", "Auto-complétion des valeurs pour les champs à choix fermé "
                  "(ex. : statut, type_commune, etat_general).", "SHOULD"),
        ("RS-04", "Prévisualisation instantanée : « N résultat(s) trouvé(s) » "
                  "avant application sur la carte.", "MUST"),
        ("RS-05", "Application de la requête = sélection des entités correspondantes "
                  "sur la carte + dans le tableau.", "MUST"),
        ("RS-06", "Historique des 10 dernières requêtes, accessible via un menu déroulant.", "COULD"),
    ],
    [1.8, 11.0, 2.0]
)

heading3(doc, "5.1.4 Onglet « Requête multicritère »")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("RM-01", "Constructeur visuel de requête : ajout de conditions par blocs "
                  "(couche, champ, opérateur, valeur).", "MUST"),
        ("RM-02", "Combinaison des conditions par ET / OU avec parenthèses logiques.", "MUST"),
        ("RM-03", "Requête spatiale : entités à moins de X mètres d'un point cliqué, "
                  "entités qui intersectent une couche cible, entités contenues dans une commune.", "SHOULD"),
        ("RM-04", "Requête combinée attributaire + spatiale (ex. : seuils EN MAUVAIS ÉTAT "
                  "situés dans un bassin versant sélectionné).", "SHOULD"),
        ("RM-05", "Sauvegarde d'une requête sous un nom (bibliothèque de requêtes utilisateur).", "COULD"),
        ("RM-06", "Affichage de la requête générée en SQL (mode expert — lecture seule).", "COULD"),
        ("RM-07", "Critère « État général d'ouvrage » : dans le constructeur de requête, "
                  "l'utilisateur peut sélectionner le champ etat_general des modèles Etat<X> "
                  "(EtatSeuil, EtatTronconSeguia, EtatBarrageRetenue, EtatKhettara, "
                  "EtatForagePuits, EtatPriseLocale, EtatMurProtection) comme condition de "
                  "filtrage, avec les 7 niveaux de l'échelle : excellent, bon, moyen-bon, "
                  "moyen, moyen-mauvais, mauvais, très mauvais. "
                  "Exemple : Seuils dont état général = 'mauvais' OU 'très mauvais'.", "MUST"),
    ],
    [1.8, 11.0, 2.0]
)

heading3(doc, "5.1.5 Onglet « Symbologie »")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("SY-01", "Mode Simple : couleur de remplissage, couleur de contour, "
                  "épaisseur de trait, opacité, taille (pour les points).", "MUST"),
        ("SY-02", "Mode Catégorisé : classification par valeur unique d'un champ "
                  "avec palette de couleurs automatique ou personnalisée.", "MUST"),
        ("SY-03", "Mode Gradué : rampe de couleur sur un champ numérique ; "
                  "choix du nombre de classes (3–9) et méthode : Jenks, Quantile, "
                  "Intervalles égaux.", "SHOULD"),
        ("SY-04", "Prévisualisation en direct dans la légende de la carte "
                  "au fur et à mesure de la modification des paramètres.", "MUST"),
        ("SY-05", "Bibliothèque personnelle : l'utilisateur sauvegarde un style sous un nom "
                  "et le retrouve pour une autre couche.", "SHOULD"),
        ("SY-06", "Import / export de style au format JSON.", "COULD"),
        ("SY-07", "Symbologie dédiée « État diagnostic » : 7 niveaux de l'échelle "
                  "excellent → très mauvais mappés sur une rampe verte → rouge.", "MUST"),
        ("SY-08", "Symbologie spéciale séguias — Largeur proportionnelle au débit : "
                  "la largeur du tracé d'un tronçon est calculée proportionnellement "
                  "au champ debit (m³/s) de TronconSeguia. Plages de largeur paramétrables "
                  "(ex. : 1 px pour ≤ 0,05 m³/s → 8 px pour ≥ 2 m³/s). "
                  "Le débit maximal des tronçons visibles sert de référence de normalisation.", "MUST"),
        ("SY-09", "Classification des couleurs des séguias par nature du matériau "
                  "(champ nature : béton, béton armé, terre, autre) : chaque valeur reçoit "
                  "une couleur distincte dans la légende. "
                  "IMPORTANT : la palette est lue dynamiquement depuis les choix disponibles "
                  "en base (NATURE_SEGUIA_CHOICES) — les valeurs NE doivent PAS être codées "
                  "en dur dans le front-end afin de rester évolutives.", "MUST"),
    ],
    [1.8, 11.0, 2.0]
)

# ─── 5.1.6 Logique de double-clic ─────────────────────────────────────────

heading3(doc, "5.1.6 Logique de double-clic (drill-down)")
body(doc, (
    "Un double-clic sur une entité déclenche un « drill-down » : "
    "la carte zoome automatiquement vers les entités filles ou associées. "
    "Cette navigation descendante suit la hiérarchie des données de la plateforme."
))
make_table(doc,
    ["Entité cliquée", "Action déclenchée", "Couche(s) affichée(s)"],
    [
        ("Province",
         "Zoom sur l'emprise de la province.\n"
         "Filtre la couche Communes pour n'afficher que celles appartenant à cette province "
         "(Province.communes — FK Commune.province).",
         "Communes (filtrées sur la province)"),
        ("Commune",
         "Zoom sur la commune.\n"
         "Filtre les Périmètres agricoles liés à cette commune "
         "(Perimetre.commune = commune sélectionnée).",
         "Périmètres (filtrés sur la commune)"),
        ("Seuil",
         "Zoom sur le bassin versant associé (Seuil.bassin_versant).\n"
         "Affiche le polygone BassinVersant.\n"
         "Affiche les tronçons ReseauHydrographique dont grid_code correspond à l'ordre "
         "de Strahler du bassin versant (classification par grid_code croissant).",
         "BassinVersant (1 polygone) +\nReseauHydrographique (filtrés par grid_code)"),
        ("Prise locale",
         "Même logique que le Seuil : zoom sur PriseLocale.bassin_versant.\n"
         "Affiche le BV + le réseau hydrographique classifié par grid_code.",
         "BassinVersant + ReseauHydrographique"),
        ("Barrage collinaire",
         "Même logique : zoom sur BarrageRetenue.bassin_versant.\n"
         "Affiche le BV + le réseau hydrographique classifié par grid_code.",
         "BassinVersant + ReseauHydrographique"),
        ("Périmètre agricole",
         "Zoom sur le périmètre.\n"
         "Active et filtre toutes les couches d'ouvrages appartenant à ce périmètre "
         "(Seuil.perimetre, Seguias.perimetre, etc.).",
         "Seuils, Séguias, Barrages, Khettaras,\nForages, Prises locales du périmètre"),
        ("BassinVersant",
         "Zoom sur le BV.\n"
         "Affiche le réseau hydrographique interne, classifié visuellement par grid_code "
         "(épaisseur croissante avec le grid_code — cours d'eau principaux plus épais).",
         "ReseauHydrographique (dans le BV)"),
    ],
    [3.5, 8.5, 6.5]
)
note(doc, "La classification du réseau hydrographique par grid_code reproduit la logique de "
          "l'ordre de Strahler calculé dans ArcGIS Pro (outil Stream Order). "
          "Plus le grid_code est élevé, plus le cours d'eau est important : "
          "afficher en trait plus épais et de couleur plus foncée.")

doc.add_page_break()

# ─── 5.2 Zone centrale ─────────────────────────────────────────────────────

heading2(doc, "5.2 Zone centrale — Trois modes de présentation")

body(doc, (
    "La zone centrale est la zone de travail principale. Elle adopte un système d'onglets "
    "offrant trois vues sur les mêmes données : Carte, Dashboard, Tableau. "
    "Les trois vues partagent le même état de sélection et sont synchronisées en permanence."
))

heading3(doc, "5.2.1 Onglet « Carte »")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("CA-01", "Rendu cartographique vectoriel interactif des couches PostGIS via GeoJSON "
                  "(ou MVT selon performance).", "MUST"),
        ("CA-02", "Fond de carte interchangeable : OpenStreetMap, sans fond (fond neutre). "
                  "Voir CA-15 pour les fonds Esri dédiés.", "SHOULD"),
        ("CA-03", "Barre d'outils flottante : zoom +/−, zoom sur emprise, outil main "
                  "(pan), mesure de distance, mesure de surface, plein écran.", "MUST"),
        ("CA-04", "Barre de statut basse : SRID actif, échelle graphique, coordonnées "
                  "du curseur, nombre d'entités sélectionnées.", "MUST"),
        ("CA-05", "Mini-carte de localisation (overview map) en coin inférieur droit.", "COULD"),
        ("CA-06", "Info-bulle riche sur survol : afficher les 5 attributs principaux "
                  "de l'entité sous le curseur.", "MUST"),
        ("CA-07", "Popup de détail sur clic : afficher tous les attributs + lien "
                  "vers la fiche détaillée Django de l'entité.", "MUST"),
        ("CA-08", "Légende dynamique affichant les symboles des couches visibles.", "MUST"),
        ("CA-09", "OUTIL D'EXPORTATION CARTOGRAPHIQUE — Choix du format : A4, A3, A2, A1, A0.", "MUST"),
        ("CA-10", "OUTIL D'EXPORTATION — Orientation : Portrait / Paysage.", "MUST"),
        ("CA-11", "OUTIL D'EXPORTATION — Résolution : 72 dpi (web), 150 dpi (brouillon), "
                  "300 dpi (impression).", "MUST"),
        ("CA-12", "OUTIL D'EXPORTATION — Éléments optionnels : Titre, Légende, "
                  "Flèche Nord, Échelle graphique, Logo, Date.", "MUST"),
        ("CA-13", "OUTIL D'EXPORTATION — Prévisualisation du layout avant export.", "SHOULD"),
        ("CA-14", "OUTIL D'EXPORTATION — Formats de sortie : PDF, PNG, SVG.", "MUST"),
        ("CA-15", "Bouton de sélection du fond de carte Esri : bouton flottant sur la carte "
                  "permettant de basculer entre trois fonds Esri :\n"
                  "  1. Esri World Imagery (satellite haute résolution)\n"
                  "  2. Esri World Topo Map (topographique avec relief ombré)\n"
                  "  3. Esri World Street Map (carte routière)\n"
                  "Les tuiles Esri sont chargées via les URL XYZ publiques d'ArcGIS Online. "
                  "Ce bouton est distinct du sélecteur de fond OSM (CA-02).", "MUST"),
    ],
    [1.8, 11.0, 2.0]
)

heading3(doc, "5.2.2 Onglet « Dashboard »")
body(doc, (
    "Le tableau de bord est lié à la couche active et à la sélection courante. "
    "Il se met à jour automatiquement lorsque la sélection change sur la carte."
))
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("DB-01", "Widget Histogramme : distribution des valeurs d'un champ numérique "
                  "(ex. : superficie irriguée des périmètres, débit des seuils).", "MUST"),
        ("DB-02", "Widget Camembert / Donut : répartition par valeur catégorielle "
                  "(ex. : état_general des seuils, type_commune, source_energie_pompage).", "MUST"),
        ("DB-03", "Widget Barres groupées : comparaison d'un indicateur par catégorie "
                  "(ex. : nombre de seuils par état, longueur de séguias par commune).", "MUST"),
        ("DB-04", "Widget KPI — Indicateur synthétique : nombre d'entités, somme, "
                  "moyenne, min/max d'un champ numérique.", "MUST"),
        ("DB-05", "Widget Carte choroplèthe miniature : colorier les communes ou "
                  "les périmètres par un indicateur agrégé.", "SHOULD"),
        ("DB-06", "Widget Série temporelle : graphique chronologique si un champ "
                  "date_diagnostic est disponible.", "COULD"),
        ("DB-07", "Disposition personnalisable : glisser-déposer des widgets "
                  "(4 emplacements par défaut).", "COULD"),
        ("DB-08", "Interaction croisée : cliquer sur un segment du graphique "
                  "sélectionne les entités correspondantes sur la carte.", "SHOULD"),
        ("DB-09", "EXPORT : PDF ou PNG du dashboard complet.", "SHOULD"),
        ("DB-10", "EXPORT : CSV/Excel des données sous-jacentes aux graphiques.", "MUST"),
    ],
    [1.8, 11.0, 2.0]
)

body(doc, "Graphiques préconfigurés selon la couche active :")
make_table(doc,
    ["Couche", "Graphiques préconfigurés"],
    [
        ("Périmètres agricoles",
         "Camembert statut juridique (melk/collectif/location/guich/habous) — "
         "Barres superficie totale vs irriguée par commune — "
         "KPI total bénéficiaires sélectionnés"),
        ("Seuils",
         "Donut état général (diagnostic_etat.etat_general) — "
         "Histogramme débit mobilisé — "
         "Barres nombre de seuils par périmètre"),
        ("Séguias / Tronçons",
         "Barres longueur totale par type (principale/secondaire/tertiaire) — "
         "KPI efficience calculée moyenne — "
         "Donut nature matériaux (béton/terre/béton armé)"),
        ("Bassins versants",
         "Scatter surface vs thalweg — "
         "Barres Q crue T10/T50/T100 par bassin — "
         "KPI surface totale sélectionnée"),
        ("Stations pluviométriques",
         "Série mensuelle précipitations (hauteur_moyenne) — "
         "Barres Pjmax par période de retour"),
    ],
    [4, 14.5]
)

heading3(doc, "5.2.3 Onglet « Tableau attributaire »")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("TA-01", "Grille de données paginée : 50 / 100 / 200 lignes par page, "
                  "navigation avant/arrière.", "MUST"),
        ("TA-02", "Colonnes redimensionnables par glisser.", "SHOULD"),
        ("TA-03", "Tri ascendant/descendant sur n'importe quelle colonne.", "MUST"),
        ("TA-04", "Filtre rapide par colonne : champ de saisie sous chaque en-tête.", "SHOULD"),
        ("TA-05", "Surlignage jaune des lignes correspondant aux entités sélectionnées sur la carte.", "MUST"),
        ("TA-06", "Clic sur une ligne dans le tableau → zoom et sélection de l'entité sur la carte.", "MUST"),
        ("TA-07", "Édition de cellule en mode inline (uniquement pour les rôles éditeur/admin), "
                  "avec validation du type de champ.", "SHOULD"),
        ("TA-08", "Pied de tableau : nombre total d'entités, nombre de lignes sélectionnées, "
                  "statistiques rapides (somme, moyenne) sur les colonnes numériques.", "SHOULD"),
        ("TA-09", "EXPORT : CSV avec en-têtes.", "MUST"),
        ("TA-10", "EXPORT : Excel (.xlsx) avec formatage de base.", "MUST"),
        ("TA-11", "EXPORT : JSON (GeoJSON si géométries incluses).", "SHOULD"),
        ("TA-12", "Option : exporter uniquement les entités sélectionnées ou toutes.", "MUST"),
        ("TA-13", "Option : choisir les colonnes à inclure dans l'export.", "SHOULD"),
    ],
    [1.8, 11.0, 2.0]
)

doc.add_page_break()

# ─── 5.3 Panneau droit ─────────────────────────────────────────────────────

heading2(doc, "5.3 Panneau droit — Outils et calculs avancés")

body(doc, (
    "Le panneau droit est un panneau vertical escamotable (largeur par défaut : 280 px). "
    "Il est inspiré du panneau Geoprocessing d'ArcGIS Pro. Il contient une barre de recherche "
    "d'outils, un accès « Favoris », un accès « Récents » et les boîtes à outils."
))

heading3(doc, "5.3.1 Boîtes à outils génériques (Analyse / Gestion / Conversion)")
body(doc, "Outils SIG transversaux, indépendants du domaine métier.")
make_table(doc,
    ["Boîte", "Outils inclus"],
    [
        ("Analyse (Analysis Tools)",
         "Tampon (Buffer) — Intersection (Intersect) — Union — Découpage (Clip) — "
         "Proximité Near (distance min) — Sélection par localisation"),
        ("Gestion des données",
         "Fusion (Merge) — Dissolution (Dissolve) — Jointure spatiale — "
         "Calculer un champ — Sélection par attribut"),
        ("Conversion",
         "Export GeoJSON — Export Shapefile (ZIP) — Export CSV attributaire"),
        ("Statistiques spatiales",
         "Résumé statistique par zone — Fréquences par catégorie — "
         "Densité de points par polygone — Heatmap de densité"),
    ],
    [4.5, 14.0]
)

heading3(doc, "5.3.2 Boîtes à outils métier (domaine irrigué)")

body(doc, (
    "Ces boîtes à outils sont spécifiques au domaine de la gestion des réseaux d'irrigation. "
    "Chaque boîte est associée à un type d'ouvrage ou d'entité de la plateforme. "
    "Toutes les boîtes métier (sauf Périmètre) incluent un outil Répéter. "
    "La boîte Périmètre inclut un Indice de priorisation (score composite)."
))

# ── BOX PÉRIMÈTRE ─────────────────────────────────────────────────────────────
heading3(doc, "Box « Périmètre »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Classification\ndéficitaire /\nexcédentaire",
         "Classe les périmètres selon leur bilan en eau (déficit ou excédent mensuel). "
         "Source : BilanBesoinRessources.resultats_bilan_normale / humide / seche. "
         "L'utilisateur choisit le type d'année (normale / humide / sèche) et le mois "
         "de référence. Résultat : couche colorée vert (excédent) / rouge (déficit) "
         "par périmètre.",
         "BilanBesoinRessources\nresultats_bilan_*"),
        ("Classification\npar culture\nou surface",
         "Classe les périmètres selon la culture dominante (culture avec le plus grand "
         "pourcentage dans Assolement) ou la superficie irriguée. "
         "Résultat : carte choroplèthe avec légende par catégorie de culture "
         "ou par tranche de superficie.",
         "Assolement.culture\nAssolement.pourcentage\nPerimetre.superficie_irriguee"),
        ("Indice de\npriorisation\n(Score périmètre)",
         "Calcule un score composite de priorité d'intervention pour chaque périmètre. "
         "PRINCIPE :\n"
         "  1. L'utilisateur assigne un coefficient (0–5) à chaque critère : "
         "superficie irriguée, nombre de bénéficiaires, superficie en bour, "
         "taux d'ouvrages dégradés (% ouvrages avec etat_general ≤ moyen-mauvais), "
         "volume déficitaire moyen.\n"
         "  2. Score = Σ(coefficient_i × valeur_normalisée_i)\n"
         "  3. Classification en N classes (3–5, au choix) par méthode Jenks ou quantile.\n"
         "  4. Carte choroplèthe résultante + export tableau des scores.\n"
         "Cliquer sur « Exécuter » déclenche le calcul côté serveur.",
         "Perimetre + EtatSeuil\n+ EtatTronconSeguia\n+ BilanBesoinRessources"),
    ],
    [2.5, 9.5, 6.5]
)

# ── BOX SEUIL ────────────────────────────────────────────────────────────────
heading3(doc, "Box « Seuil »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Score ouvrage /\nIndice de priorité",
         "Calcule un score de dégradation pour chaque seuil selon les critères du "
         "diagnostic structuré (EtatSeuil). L'utilisateur assigne un coefficient (0–5) "
         "à chacun des 10 critères : etat_structurel_digue, affouillement_aval, "
         "envasement_retenue, murs_guideaux, radier_aval, etat_vannes, dessableur, "
         "degradation_beton, infiltration_fuite, limiteur_debit.\n"
         "Score = Σ(coefficient_i × note_i) / Σ(coefficients_max).\n"
         "Classification et carte colorée de priorité d'intervention.",
         "EtatSeuil (notes 0–5)"),
        ("Répéter",
         "Applique l'analyse en cours (score, requête, export) à tous les seuils "
         "sélectionnés en mode batch. Barre de progression + log par seuil.",
         "Seuil (sélection active)"),
        ("Statistiques\ndébit",
         "Histogramme des débits mobilisés (l/s) des seuils sélectionnés + "
         "statistiques (min, max, moyenne, médiane).",
         "Seuil.debit_mobilise"),
        ("Carte efficience\nréseau",
         "Carte choroplèthe des seuils selon leur efficience_reseaux (0–1). "
         "Rampe rouge (faible) → vert (élevée).",
         "Seuil.efficience_reseaux"),
    ],
    [2.5, 9.5, 6.5]
)

# ── BOX SÉGUIA ───────────────────────────────────────────────────────────────
heading3(doc, "Box « Séguia / Tronçon »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Efficience\ndes tronçons\n(PI + PV)",
         "Déclenche le recalcul de l'efficience hydraulique pour les tronçons sélectionnés "
         "en appelant les fonctions déjà définies dans la plateforme :\n"
         "  • PI = Perte par Infiltration (m³/s) selon la nature et les dimensions du tronçon\n"
         "  • PV = Perte par Évaporation (m³/s) selon la surface mouillée et l'ET0\n"
         "  • Efficience calculée = (Débit entrée − PI − PV) / Débit entrée × 100 (%)\n"
         "Résultats écrits dans TronconSeguia.efficience_calculee, "
         ".perte_infiltration_m3s, .perte_vaporisation_m3s.",
         "TronconSeguia\n+ Perimetre.et0_mm_jour\n(fonctions existantes)"),
        ("Débit Manning\ndu tronçon",
         "Calcule le débit hydraulique du tronçon par la formule de Manning-Strickler :\n"
         "  Q = (1/n) × A × R^(2/3) × S^(1/2)\n"
         "Paramètres lus depuis TronconSeguia : forme, longueur, largeur_meroire, "
         "hauteur_eau, fruit_de_berge, epaisseur_parois, diametre.\n"
         "n (coefficient Manning) est saisi par l'utilisateur ou lu depuis une valeur "
         "par défaut selon la nature du matériau.\n"
         "Résultat affiché dans un panel + colonne calculée dans le tableau.",
         "TronconSeguia\n(formule déjà définie\ndans la plateforme)"),
        ("Score ouvrage /\nIndice de priorité",
         "Même logique que le Seuil : l'utilisateur assigne des coefficients aux 7 critères "
         "de EtatTronconSeguia (fissures_revetement, infiltration_fuite, obstructions_debris, "
         "erosion_berges, sedimentation_fond, ouvrages_regulation, spalling_beton). "
         "Score composite → classification → carte.",
         "EtatTronconSeguia (notes 0–5)"),
        ("Répéter",
         "Applique l'analyse (efficience, Manning, score) à tous les tronçons "
         "sélectionnés en batch.",
         "TronconSeguia (sélection)"),
    ],
    [2.5, 9.5, 6.5]
)

# ── BOX PRISE LOCALE ─────────────────────────────────────────────────────────
heading3(doc, "Box « Prise Locale »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Score ouvrage /\nIndice de priorité",
         "Coefficients sur les 5 critères de EtatPriseLocale : "
         "envasement_sedimentation_entree, degradation_revetement, "
         "accumulation_debris_vegetation, etat_dispositifs_regulation, "
         "protection_crues_debordements. "
         "Score → classification → carte.",
         "EtatPriseLocale (notes 0–5)"),
        ("Répéter",
         "Batch sur toutes les prises locales sélectionnées.",
         "PriseLocale (sélection)"),
        ("Débit dérivé\nvs besoin",
         "Compare le debit_derive de la prise locale "
         "avec le besoin en eau du périmètre associé "
         "(résultat BilanBesoinRessources). Graphique en barres par prise.",
         "PriseLocale.debit_derive\n+ BilanBesoinRessources"),
    ],
    [2.5, 9.5, 6.5]
)

# ── BOX BARRAGE COLLINAIRE ────────────────────────────────────────────────────
heading3(doc, "Box « Barrage Collinaire »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Score ouvrage /\nIndice de priorité",
         "Coefficients sur les 4 critères de EtatBarrageRetenue : "
         "affouillement_pied_digue_aval, taux_envasement_retenue, "
         "regulation_debits_aval, fonctionnement_ouvrages_prise_eau. "
         "Score → classification → carte.",
         "EtatBarrageRetenue (notes 0–5)"),
        ("Répéter",
         "Batch sur tous les barrages collinaires sélectionnés.",
         "BarrageRetenue (sélection)"),
        ("Bilan volume\napports vs besoins",
         "Compare les apports mensuels du barrage (BilanOuvrageAssocie.apports_mensuels_*) "
         "avec les besoins du périmètre associé. "
         "Graphique mensuel en barres empilées pour les 3 types d'années.",
         "BilanOuvrageAssocie\n+ BilanBesoinRessources"),
    ],
    [2.5, 9.5, 6.5]
)

# ── BOX PUITS / FORAGE ────────────────────────────────────────────────────────
heading3(doc, "Box « Puits / Forage »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Score ouvrage /\nIndice de priorité",
         "Coefficients sur les 4 critères de EtatForagePuits : "
         "qualite_physico_chimique_eau, degradation_structurelle_forage, "
         "colmatage_forage, etat_equipements. "
         "Score → classification → carte.",
         "EtatForagePuits (notes 0–5)"),
        ("Répéter",
         "Batch sur tous les forages/puits sélectionnés.",
         "ForagePuits (sélection)"),
        ("Comparaison\npar source\nd'énergie",
         "Carte thématique et graphique en camembert des forages "
         "selon leur source_energie_pompage (réseau électrique, solaire, diesel, hybride).",
         "ForagePuits.source_energie_pompage"),
    ],
    [2.5, 9.5, 6.5]
)

# ── BOX KHETTARAT ─────────────────────────────────────────────────────────────
heading3(doc, "Box « Khettarat »")
make_table(doc,
    ["Outil", "Description", "Source de données"],
    [
        ("Score ouvrage /\nIndice de priorité",
         "Coefficients sur les 4 critères de EtatKhettara : "
         "envasement_ensablement_fond, degradation_beton, "
         "accessibilite_entretien, stabilite_galerie_principale. "
         "Score → classification → carte.",
         "EtatKhettara (notes 0–5)"),
        ("Répéter",
         "Batch sur toutes les khettaras sélectionnées.",
         "Khettara (sélection)"),
        ("Carte débit\nvs longueur",
         "Scatter plot débit (m³/s) vs longueur (m) + carte proportionnelle "
         "aux débits (cercles de taille proportionnelle à Khettara.debit).",
         "Khettara.debit\nKhettara.longueur"),
    ],
    [2.5, 9.5, 6.5]
)

body(doc, "")
note(doc, "Outil Répéter — Principe commun à toutes les boîtes métier (sauf Périmètre) : "
          "l'utilisateur sélectionne une ou plusieurs entités sur la carte, configure "
          "l'outil, puis clique « Répéter » pour appliquer le même calcul à toutes les "
          "entités sélectionnées en batch. Une barre de progression et un log d'exécution "
          "sont affichés. Les résultats sont présentés dans un tableau comparatif et "
          "optionnellement exportables en Excel.")

heading3(doc, "5.3.3 Exigences transversales sur les outils")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("OUT-01", "Chaque outil s'ouvre dans un formulaire modal latéral avec : "
                   "description courte, aide contextuelle par champ (info-bulle ?).", "MUST"),
        ("OUT-02", "Barre de progression pendant l'exécution d'un outil long.", "SHOULD"),
        ("OUT-03", "Log d'exécution consultable (messages, durée, erreurs).", "SHOULD"),
        ("OUT-04", "Le résultat d'un outil est automatiquement ajouté dans le gestionnaire "
                   "de couches avec un nom par défaut modifiable.", "MUST"),
        ("OUT-05", "Historique des 5 derniers outils utilisés (avec paramètres mémorisés), "
                   "affiché dans la section « Récents ».", "SHOULD"),
        ("OUT-06", "Système de favoris : l'utilisateur épingle ses outils fréquents "
                   "en haut du panneau.", "SHOULD"),
        ("OUT-07", "Barre de recherche en temps réel sur le nom des outils.", "MUST"),
    ],
    [1.8, 11.0, 2.0]
)

heading3(doc, "5.3.4 Calculatrice de champ")
make_table(doc,
    ["ID Req.", "Description", "Priorité"],
    [
        ("CF-01", "Éditeur d'expression multi-lignes avec coloration syntaxique.", "SHOULD"),
        ("CF-02", "Fonctions disponibles : arithmétiques, texte (upper, lower, concat), "
                  "date, géométriques (ST_Length, ST_Area, ST_Centroid).", "SHOULD"),
        ("CF-03", "Aperçu de la valeur calculée sur un échantillon de 5 entités "
                  "avant application.", "SHOULD"),
        ("CF-04", "Résultat écrit dans un nouveau champ calculé (non persistant en base "
                  "par défaut, sauf action explicite de l'éditeur).", "SHOULD"),
    ],
    [1.8, 11.0, 2.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  6. INTERACTIONS ET SYNCHRONISATION
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "6. Interactions et synchronisation entre vues")

body(doc, (
    "L'état de l'application est centré sur deux variables partagées entre les trois vues :"
))
bullet(doc, "couche_active : la couche couramment interrogée (détermine les champs disponibles "
            "dans les requêtes, le tableau, les graphiques).")
bullet(doc, "selection_active : l'ensemble des identifiants (pk Django) des entités sélectionnées "
            "dans la couche active.")

make_table(doc,
    ["Événement source", "Vue source", "Effet sur les autres vues"],
    [
        ("Sélection d'entités via outil de sélection",
         "Carte",
         "Tableau : surligne les lignes correspondantes et scroll vers la première.\n"
         "Dashboard : recalcule tous les graphiques sur la sélection courante."),
        ("Clic sur une ligne du tableau",
         "Tableau",
         "Carte : zoom et surligne l'entité.\n"
         "Dashboard : met à jour avec entité unique si applicable."),
        ("Clic sur un segment de graphique",
         "Dashboard",
         "Carte : sélectionne et surligne les entités du segment.\n"
         "Tableau : filtre sur les lignes du segment."),
        ("Application d'une requête attributaire",
         "Panneau gauche",
         "Carte, Tableau, Dashboard : tous recalculent sur le nouveau jeu résultat."),
        ("Changement de couche active",
         "Panneau gauche",
         "Tableau : affiche les attributs de la nouvelle couche (sélection réinitialisée).\n"
         "Dashboard : affiche les graphiques préconfigurés de la nouvelle couche."),
        ("Résultat d'un outil (nouvelle couche)",
         "Panneau droit",
         "Gestionnaire de couches : ajoute la couche résultat.\n"
         "Carte : affiche la nouvelle couche automatiquement."),
    ],
    [4.5, 3.5, 10.5]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  7. INTERFACES DE PROGRAMMATION (API)
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "7. Interfaces de programmation (API Django)")

body(doc, (
    "Le module Carte requiert de nouvelles vues Django exposant des endpoints JSON/GeoJSON "
    "sous le préfixe /carte/api/. Ces endpoints sont consommés par le front-end JavaScript "
    "de la carte interactive."
))

heading2(doc, "7.1 Endpoints GeoJSON des couches")
make_table(doc,
    ["URL", "Méthode", "Paramètres", "Description"],
    [
        ("/carte/api/couche/<nom>/", "GET",
         "bbox, srid, limit, offset, fields",
         "Retourne les entités GeoJSON de la couche <nom>. "
         "Pagination par bbox et offset/limit."),
        ("/carte/api/couche/<nom>/<pk>/", "GET",
         "—",
         "Retourne une entité GeoJSON complète avec tous ses attributs."),
        ("/carte/api/couche/<nom>/extent/", "GET",
         "—",
         "Retourne la bounding box (xmin, ymin, xmax, ymax) de la couche."),
        ("/carte/api/couches/", "GET",
         "—",
         "Liste des couches disponibles avec métadonnées (nom, type géom, champs, groupe)."),
    ],
    [5.0, 2.0, 3.5, 8.0]
)

heading2(doc, "7.2 Endpoints de requête et sélection")
make_table(doc,
    ["URL", "Méthode", "Corps / Paramètres", "Description"],
    [
        ("/carte/api/requete/simple/", "POST",
         "{couche, champ, operateur, valeur}",
         "Retourne les pk des entités correspondantes."),
        ("/carte/api/requete/multicritere/", "POST",
         "{couche, conditions: [...], logique}",
         "Retourne les pk des entités correspondant à l'expression logique."),
        ("/carte/api/requete/spatiale/", "POST",
         "{couche, type_spatial, geometrie_ref, distance_m}",
         "Requête spatiale : intersects, within, dwithin."),
        ("/carte/api/couche/<nom>/champs/", "GET",
         "—",
         "Liste des champs de la couche avec type, verbose_name et indicateur à choix fermé."),
        ("/carte/api/couche/<nom>/champs/<champ>/valeurs/", "GET",
         "—",
         "Valeurs distinctes d'un champ (pour choix fermés et filtres). "
         "Retourne les valeurs réelles + libellés depuis les CHOICES Django. "
         "REQUIS pour l'évolutivité — jamais coder ces valeurs en front-end."),
    ],
    [5.0, 2.0, 4.5, 7.0]
)

heading2(doc, "7.3 Endpoints des outils géospatiaux")
make_table(doc,
    ["URL", "Méthode", "Description"],
    [
        ("/carte/api/outils/buffer/",       "POST", "Génère un tampon autour d'une couche ou d'une sélection."),
        ("/carte/api/outils/intersection/", "POST", "Intersection entre deux couches."),
        ("/carte/api/outils/union/",        "POST", "Union de deux couches."),
        ("/carte/api/outils/dissolve/",     "POST", "Dissolution d'une couche par champ."),
        ("/carte/api/outils/near/",         "POST", "Distance minimale entre deux couches."),
        ("/carte/api/outils/stats/",        "POST", "Statistiques par zone (count, sum, avg, min, max)."),
    ],
    [5.5, 2.0, 11.0]
)

heading2(doc, "7.4 Endpoints d'export")
make_table(doc,
    ["URL", "Méthode", "Description"],
    [
        ("/carte/api/export/csv/",      "POST", "Export CSV d'une couche ou sélection."),
        ("/carte/api/export/excel/",    "POST", "Export Excel (.xlsx) d'une couche ou sélection."),
        ("/carte/api/export/geojson/",  "POST", "Export GeoJSON d'une couche ou sélection."),
        ("/carte/api/export/carte/",    "POST", "Export cartographique : PDF/PNG, format A4–A0, DPI, éléments."),
        ("/carte/api/export/dashboard/","POST", "Export dashboard en PDF ou PNG."),
    ],
    [5.5, 2.0, 11.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  8. SÉCURITÉ ET DROITS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "8. Sécurité et gestion des droits")

heading2(doc, "8.1 Matrice des rôles")
body(doc, (
    "Le module s'appuie sur le modèle Utilisateur existant (compte.Utilisateur) "
    "avec trois rôles : visiteur, opérateur, éditeur."
))
make_table(doc,
    ["Fonctionnalité", "visiteur", "opérateur", "éditeur"],
    [
        ("Consultation carte (toutes couches)", "OUI", "OUI", "OUI"),
        ("Requête simple et sélection", "NON", "OUI", "OUI"),
        ("Requête multicritère et requête spatiale", "NON", "OUI", "OUI"),
        ("Outils d'analyse (tampon, intersection, etc.)", "NON", "OUI", "OUI"),
        ("Export CSV / Excel / GeoJSON", "NON", "OUI", "OUI"),
        ("Export cartographique PDF/PNG", "NON", "OUI", "OUI"),
        ("Sauvegarde de styles symbologie", "NON", "OUI", "OUI"),
        ("Sauvegarde de requêtes nommées", "NON", "OUI", "OUI"),
        ("Édition de cellule dans le tableau attributaire", "NON", "NON", "OUI"),
        ("Création de couches issues d'outils (persistance)", "NON", "NON", "OUI"),
        ("Accès à l'administration Django (/admin)", "NON", "NON", "OUI"),
    ],
    [7.0, 2.0, 2.3, 2.3]
)

heading2(doc, "8.2 Exigences de sécurité")
make_table(doc,
    ["ID", "Exigence"],
    [
        ("SEC-01", "Tous les endpoints /carte/api/* doivent exiger @login_required. "
                   "Retourner HTTP 403 pour accès non autorisé."),
        ("SEC-02", "Ajouter @role_required (décorateur à créer dans compte/) sur les vues "
                   "nécessitant opérateur ou éditeur — combler le gap de sécurité mentionné "
                   "dans CLAUDE.md."),
        ("SEC-03", "Protection CSRF sur tous les endpoints POST/PUT/DELETE."),
        ("SEC-04", "Validation côté serveur de tous les paramètres GeoJSON entrants "
                   "(format, SRID, taille max 5 Mo)."),
        ("SEC-05", "Les requêtes SQL avancées (mode expert) doivent passer par l'ORM Django "
                   "ou une liste blanche de fonctions PostGIS — pas de SQL brut non filtré."),
        ("SEC-06", "Les fichiers exportés sont générés en mémoire et renvoyés en streaming "
                   "(pas de fichiers temporaires persistants sur disque)."),
    ],
    [1.5, 17.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  9. EXIGENCES NON FONCTIONNELLES
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "9. Exigences non fonctionnelles")

heading2(doc, "9.1 Performance")
make_table(doc,
    ["ID", "Exigence", "Cible"],
    [
        ("PERF-01", "Temps de chargement initial d'une couche légère (< 500 entités)", "< 2 s"),
        ("PERF-02", "Temps de chargement d'une couche lourde (séguias, réseau hydro)", "< 5 s"),
        ("PERF-03", "Réponse à une requête attributaire simple", "< 1 s"),
        ("PERF-04", "Réponse à un outil tampon sur 100 entités", "< 3 s"),
        ("PERF-05", "Génération d'un export PDF A4 300 dpi", "< 10 s"),
        ("PERF-06", "Rendu carte après zoom/pan (tiles ou GeoJSON re-fetch)", "< 1 s"),
    ],
    [1.5, 11.5, 2.5]
)
note(doc, "Pour les couches volumineuses (réseau hydrographique > 10 000 tronçons), "
          "utiliser MVT (Mapbox Vector Tiles) via pg_tileserv pour éviter les transferts GeoJSON "
          "trop lourds. Cette décision est à confirmer en phase de prototypage.")

heading2(doc, "9.2 Fiabilité et disponibilité")
bullet(doc, "Les opérations de lecture (affichage, requêtes) ne doivent jamais modifier la base de données.")
bullet(doc, "En cas d'erreur d'un outil géospatial, la carte doit rester fonctionnelle "
            "(pas de crash global de la page).")
bullet(doc, "Un message d'erreur explicite doit être affiché à l'utilisateur "
            "(toast notification + log consultable).")

heading2(doc, "9.3 Compatibilité navigateurs")
bullet(doc, "Chrome 120+, Firefox 120+, Edge 120+ (navigateurs modernes supportant WebGL).")
bullet(doc, "Résolution minimale : 1280 × 720 px.")
bullet(doc, "Interface responsive : utilisable sur tablette (1024 px), panneaux latéraux rétractables.")

heading2(doc, "9.4 Maintenabilité")
bullet(doc, "Le module « carte » est une nouvelle application Django (carte/) qui complète "
            "les applications existantes — ne pas modifier les modèles des autres apps.")
bullet(doc, "Les couches sont configurées déclarativement dans un registre Python "
            "(LAYER_REGISTRY dans carte/layers.py) — ajouter une couche = une ligne dans ce registre.")
bullet(doc, "Les outils géospatiaux utilisent GDAL/GEOS via GeoDjango, "
            "pas de dépendances externes supplémentaires.")
bullet(doc, "Code commenté en français, conforme aux conventions Django existantes.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  10. ARCHITECTURE TECHNIQUE CIBLE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "10. Architecture technique cible")

heading2(doc, "10.1 Schéma global")
body(doc, "L'architecture suit le pattern Django MVT enrichi d'une couche API JSON :")

layers_arch = [
    ("Navigateur (front-end)",
     "MapLibre GL JS (rendu vectoriel), Chart.js (graphiques), "
     "JavaScript vanilla ou Alpine.js (interactions panneau), "
     "Fetch API (appels AJAX vers /carte/api/*)."),
    ("Django Views (/carte/)",
     "Vue principale HTML : carte/views.py → render('carte/index.html')\n"
     "API GeoJSON : carte/api_views.py → JsonResponse / StreamingHttpResponse\n"
     "Outils : carte/tools.py → wrappeurs GDAL/GeoDjango"),
    ("ORM GeoDjango",
     "Accès PostGIS via les QuerySets GIS : .filter(geometrie__intersects=...), "
     ".annotate(distance=...), .transform(srid=...)."),
    ("PostGIS",
     "Tables des 5 apps Django + fonctions spatiales (ST_Buffer, ST_Intersection, "
     "ST_Distance, ST_DWithin, ST_Simplify pour les tuiles)."),
]
for layer, desc in layers_arch:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(f"[ {layer} ]  ")
    run1.bold = True
    run1.font.color.rgb = C_BLEU_H2
    run1.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

heading2(doc, "10.2 Structure de fichiers proposée")
code_block = (
    "plateformeSIG/\n"
    "  carte/\n"
    "    models.py          ← Province, Commune (existants) + StyleCouche, RequeteNommee (nouveaux)\n"
    "    layers.py          ← LAYER_REGISTRY : déclaration de toutes les couches\n"
    "    views.py           ← Vue HTML principale\n"
    "    api_views.py       ← Endpoints GeoJSON, requêtes, export\n"
    "    tools.py           ← Logique des outils géospatiaux (tampon, etc.)\n"
    "    serializers.py     ← Sérialisation GeoJSON par couche\n"
    "    urls.py            ← Routes /carte/ et /carte/api/*\n"
    "    templates/carte/\n"
    "      index.html       ← Page principale du module carte\n"
    "      partials/\n"
    "        panel_left.html\n"
    "        panel_right.html\n"
    "        toolbar_map.html\n"
    "    static/carte/\n"
    "      js/\n"
    "        map.js         ← Initialisation MapLibre GL\n"
    "        layers.js      ← Chargement GeoJSON dynamique\n"
    "        query.js       ← Requêtes attributaires\n"
    "        tools.js       ← Appels API outils\n"
    "        dashboard.js   ← Chart.js widgets\n"
    "        table.js       ← Tableau attributaire\n"
    "      css/\n"
    "        carte.css      ← Styles panneaux et carte\n"
)
p = doc.add_paragraph(code_block)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
p.paragraph_format.left_indent = Cm(0.5)

heading2(doc, "10.3 Nouveaux modèles Django à créer")
make_table(doc,
    ["Modèle", "Champs principaux", "Rôle"],
    [
        ("StyleCouche",
         "utilisateur (FK), nom_couche, nom_style, parametres (JSONField), created_at",
         "Bibliothèque de styles personnels de l'utilisateur."),
        ("RequeteNommee",
         "utilisateur (FK), nom, couche, expression (JSONField), created_at",
         "Sauvegarde des requêtes multicritères nommées."),
    ],
    [3.5, 9.0, 6.0]
)
note(doc, "Ces deux modèles sont rattachés à l'app carte/ et nécessitent une migration Django.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  11. CONTRAINTES ET DÉPENDANCES
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "11. Contraintes et dépendances")

heading2(doc, "11.1 Contraintes techniques")
make_table(doc,
    ["Contrainte", "Impact"],
    [
        ("GDAL 3.12 (OSGeo4W) déjà installé",
         "Les opérations vectorielles doivent utiliser GDAL via GeoDjango. "
         "Pas d'ajout de dépendance GIS externe (Shapely, Fiona, Geopandas) "
         "sans validation préalable."),
        ("SRID 4326 en base",
         "Tous les GeoJSON retournés seront en SRID 4326. "
         "Les calculs de distance/surface nécessitent une re-projection "
         "en EPSG:26191 (Nord Maroc) côté serveur."),
        ("Django 6 — pas de DRF",
         "Les API JSON sont implémentées avec des vues Django standard "
         "(JsonResponse), pas Django REST Framework."),
        ("Calendrier hydrologique Sep→Aoû",
         "Les graphiques de séries temporelles mensuelles doivent afficher "
         "les 12 mois dans l'ordre Sep→Aoû, pas Jan→Déc."),
        ("Champ statut sur ouvrages",
         "Les couches Diagnostic ont un champ statut (non_valide / valide). "
         "La symbologie par défaut doit distinguer visuellement ces deux états."),
    ],
    [6.0, 12.5]
)

heading2(doc, "11.2 Dépendances avec les autres modules")
make_table(doc,
    ["Module dépendant", "Nature de la dépendance"],
    [
        ("diagnostic",
         "Fournit les 7 types d'ouvrages hydrauliques + périmètres. "
         "Le module Carte ne doit pas modifier leurs modèles."),
        ("analyse_hydrologique",
         "Fournit bassins versants, stations, réseau hydrographique. "
         "Les résultats de crues (ResultatAnalyseHydrologique) peuvent "
         "être affichés dans le dashboard."),
        ("Besions_Ressources",
         "Fournit les stations climatiques et les résultats de bilan. "
         "Les bilans mensuels peuvent alimenter des graphiques dans le dashboard."),
        ("compte",
         "Fournit le modèle Utilisateur et les rôles. "
         "Les droits d'accès du module Carte s'appuient sur compte.Utilisateur.role."),
        ("efficiences (futur)",
         "Quand ce module sera développé, ses couches (séguias avec efficience) "
         "seront automatiquement disponibles si elles ont un champ geometrie "
         "et sont enregistrées dans LAYER_REGISTRY."),
    ],
    [4.5, 14.0]
)

heading2(doc, "11.3 Exigence fondamentale — Durabilité et Évolutivité")
body(doc, (
    "POINT CRITIQUE : la base de données utilise des listes de choix (CHOICES) pour de "
    "nombreux champs (nature des matériaux, type de seuil, source d'énergie, état général…). "
    "Ces listes SONT ET SERONT modifiées au fil du temps — de nouvelles valeurs peuvent "
    "être ajoutées sans toucher au schéma Django (via migration ou directement en base)."
))
body(doc, (
    "Le module Carte doit être conçu pour que l'ajout d'une nouvelle valeur dans une "
    "liste de choix n'implique AUCUNE modification du code front-end. "
    "Cette contrainte s'applique à tous les composants qui font référence à ces valeurs."
))
make_table(doc,
    ["Composant", "Règle d'implémentation", "Exemple"],
    [
        ("Symbologie\n(classification)",
         "La liste des valeurs pour la classification catégorielle ou graduée "
         "doit être lue dynamiquement depuis l'API Django "
         "(/carte/api/couche/<nom>/champs/<champ>/valeurs/), "
         "jamais codée en dur dans le JavaScript.",
         "nature de séguia : ['béton', 'terre', …]\n"
         "Si 'PVC' est ajouté en base, la légende l'affiche automatiquement."),
        ("Filtres du tableau\nattributaire",
         "Les valeurs proposées dans les filtres déroulants d'une colonne "
         "doivent être alimentées par une requête GROUP BY côté serveur, "
         "pas par une liste statique.",
         "source_energie_pompage : dropdown alimenté\npar les valeurs distinctes en base"),
        ("Requête simple\net multicritère",
         "Les valeurs proposées en auto-complétion pour les champs à choix fermé "
         "doivent être récupérées depuis l'endpoint "
         "/carte/api/couche/<nom>/champs/<champ>/valeurs/.",
         "etat_general : liste lue depuis\nEAT_CONSTRUCTION_DIAG_CHOICES côté serveur"),
        ("Dashboard\n(graphiques)",
         "Les catégories des camemberts / histogrammes doivent être construites "
         "dynamiquement à partir des valeurs distinctes retournées par l'API. "
         "Pas de mapping code couleur figé par valeur.",
         "Donut nature séguias : nombre de segments\nvariable selon valeurs en base"),
        ("Boîtes à outils\nmétier",
         "Les critères de scoring des outils Score ouvrage doivent lire "
         "les champs et leur libellé depuis l'API (introspection du modèle Etat<X>), "
         "pas depuis une liste figée dans le front-end.",
         "Si un nouveau critère est ajouté à EtatSeuil,\nil apparaît automatiquement\n"
         "dans le formulaire de scoring"),
        ("LAYER_REGISTRY\n(back-end)",
         "Toute modification de la liste des choix d'un champ est transparente pour "
         "le registre de couches : le registre déclare les champs et leur type, "
         "mais pas les valeurs permises.",
         "Ajout d'un nouveau type de seuil :\naucune modification dans layers.py"),
    ],
    [2.8, 8.5, 7.2]
)
note(doc, "Cette exigence d'évolutivité est le CRITÈRE DE CONCEPTION LE PLUS IMPORTANT "
          "du module Carte. Elle conditionne l'architecture de l'API et du front-end. "
          "Un composant qui lit ses valeurs depuis le serveur à chaque affichage est "
          "toujours préférable à un composant qui les a en dur, même si c'est légèrement "
          "plus coûteux en requêtes.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  12. CRITÈRES D'ACCEPTATION
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "12. Critères d'acceptation")

body(doc, (
    "Les critères d'acceptation ci-dessous servent de base aux tests de recette "
    "à réaliser avant la mise en production du module."
))

make_table(doc,
    ["ID", "Scénario de test", "Résultat attendu"],
    [
        ("CA-01",
         "Ouvrir /carte/ avec un compte visiteur",
         "La carte s'affiche avec les couches visibles en lecture seule. "
         "Le panneau gauche montre uniquement les onglets Couches et Symbologie. "
         "Pas de bouton Requête, Outils, Export."),
        ("CA-02",
         "Ouvrir /carte/ avec un compte opérateur",
         "Tous les panneaux et onglets sont accessibles. "
         "Le tableau attributaire n'affiche pas de bouton Éditer."),
        ("CA-03",
         "Activer/désactiver la visibilité de la couche Périmètres",
         "La couche disparaît/réapparaît sur la carte sans rechargement de page."),
        ("CA-04",
         "Sélectionner 3 seuils par rectangle",
         "3 entités sont surlignées sur la carte + 3 lignes surlignées dans le tableau "
         "+ le dashboard recalcule sur ces 3 seuils."),
        ("CA-05",
         "Requête simple : Seuils dont statut = 'valide'",
         "Seuls les seuils validés sont sélectionnés sur la carte. "
         "Le compteur affiche le bon nombre."),
        ("CA-06",
         "Appliquer un tampon de 500 m sur un périmètre sélectionné",
         "Une nouvelle couche temporaire apparaît dans le gestionnaire de couches "
         "avec la géométrie du tampon."),
        ("CA-07",
         "Exporter le tableau attributaire des seuils en Excel",
         "Un fichier .xlsx est téléchargé avec toutes les colonnes et les données correctes."),
        ("CA-08",
         "Exporter la carte en PDF A3 Paysage 300 dpi",
         "Un fichier PDF est généré avec l'étendue visible, la légende, "
         "la flèche nord et l'échelle."),
        ("CA-09",
         "Clic sur une barre du dashboard (Donut état des seuils, segment 'Mauvais')",
         "Les seuils de cet état sont sélectionnés sur la carte et le tableau."),
        ("CA-10",
         "Accéder à /carte/api/couche/seuils/ sans être connecté",
         "Réponse HTTP 403 (Forbidden)."),
    ],
    [1.5, 7.0, 10.0]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  13. LIVRABLES ATTENDUS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "13. Livrables attendus")

make_table(doc,
    ["#", "Livrable", "Format", "Description"],
    [
        ("L-01", "Code source du module Carte",
         "Python / HTML / JS / CSS",
         "Application Django carte/ complète : models, views, api_views, tools, "
         "serializers, urls, templates, static."),
        ("L-02", "Migrations Django",
         "Fichiers Python",
         "Migrations pour les modèles StyleCouche et RequeteNommee."),
        ("L-03", "Registre de couches",
         "carte/layers.py",
         "Déclaration de toutes les couches disponibles avec métadonnées."),
        ("L-04", "Tests unitaires et d'intégration",
         "carte/tests.py",
         "Tests couvrant les API GeoJSON, les outils et les exports. "
         "Couverture minimale : 70 %."),
        ("L-05", "Documentation développeur",
         "Mise à jour CLAUDE.md",
         "Section carte/ : architecture, LAYER_REGISTRY, ajout d'une couche, "
         "ajout d'un outil."),
        ("L-06", "Guide utilisateur",
         "PDF ou page HTML",
         "Guide d'utilisation des fonctionnalités du module Carte "
         "par rôle (opérateur / éditeur)."),
        ("L-07", "Rapport de recette",
         "Word / PDF",
         "Résultats des 10 critères d'acceptation avec captures d'écran."),
    ],
    [0.8, 3.5, 3.0, 11.5]
)

doc.add_paragraph()

heading2(doc, "Planning indicatif")
make_table(doc,
    ["Phase", "Contenu", "Durée estimée"],
    [
        ("Phase 0 — Prototypage",
         "Mise en place MapLibre GL, affichage d'une couche GeoJSON, "
         "API /carte/api/couches/, maquette UI.",
         "1 semaine"),
        ("Phase 1 — Couches et visualisation",
         "LAYER_REGISTRY complet, gestionnaire de couches, "
         "symbologie simple et catégorisée, fond de carte.",
         "2 semaines"),
        ("Phase 2 — Sélection et requêtes",
         "Outils de sélection, requête simple, requête multicritère, "
         "synchronisation carte ↔ tableau.",
         "2 semaines"),
        ("Phase 3 — Tableau et dashboard",
         "Tableau attributaire complet, dashboard avec 4 widgets préconfigurés, "
         "synchronisation tableau ↔ dashboard.",
         "2 semaines"),
        ("Phase 4 — Outils et exports",
         "Boîte à outils (tampon, intersection, dissolve, stats), "
         "export carte PDF, export CSV/Excel.",
         "2 semaines"),
        ("Phase 5 — Finitions et recette",
         "Sécurité (décorateurs rôles), tests, documentation, recette.",
         "1 semaine"),
    ],
    [3.5, 10.5, 2.5]
)

doc.add_paragraph()
body(doc, "Durée totale estimée : 10 semaines pour une implémentation complète de la V1 (MUST + SHOULD).")

doc.add_page_break()

# ── Pied de document ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run(
    "Document produit dans le cadre du Projet de Fin d'Études (PFE) — "
    "Plateforme SIG Ressources en Eau — Tafilalet / Midelt\n"
    "Version 1.0 — Juin 2026"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
run.font.italic = True

# ─────────────────────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"Document généré : {OUTPUT}")
